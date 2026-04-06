"""Pytest tests for the DOCX v3 report generator (generate_docx_v3_report)."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from engines.report_generation.docx_generator import generate_docx_v3_report
from engines.report_generation.v3.docx_lint import lint_docx_render_contract
from engines.report_generation.v3.render_contract import LintSeverity


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _read_docx_text(path: Path) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Shared fixture: full report dict used by most tests
# ---------------------------------------------------------------------------

_FULL_REPORT: dict = {
    "report_id": "rpt-001",
    "case_id": "case-001",
    "run_id": "run-001",
    "perspective": "neutral",
    "layer1": {
        "cover_summary": {
            "neutral_conclusion": "案件结论摘要",
            "winning_move": "关键证据决定胜负",
            "blocking_conditions": ["条件1", "条件2"],
        },
        "timeline": [
            {
                "date": "2024-01-01",
                "event": "签订借贷协议",
                "source": "ev-001",
                "disputed": False,
                "tag": "事实",
            },
        ],
        "evidence_priorities": [
            {
                "evidence_id": "ev-001",
                "title": "借款合同",
                "priority": "核心证据",
                "reason": "直接证明借贷关系",
            },
        ],
        "evidence_traffic_lights": [],
        "scenario_tree_summary": "",
    },
    "layer2": {
        "fact_base": [
            {
                "fact_id": "f-001",
                "description": "2024年1月1日双方签订借款合同",
                "source_evidence_ids": ["ev-001"],
                "tag": "事实",
            },
        ],
        "issue_map": [
            {
                "issue_id": "issue-001",
                "issue_title": "借款关系是否成立",
                "depth": 0,
                "plaintiff_thesis": "原告已出借款项",
                "defendant_thesis": "被告否认收款",
                "decisive_evidence": ["ev-001"],
                "current_gaps": ["转账记录缺失"],
                "outcome_sensitivity": "高",
                "tag": "推断",
            },
        ],
        "evidence_cards": [
            {
                "evidence_id": "ev-001",
                "q1_what": "借款合同原件",
                "q2_target": "证明借贷关系成立",
                "q3_key_risk": "被告质疑真实性",
                "q4_best_attack": "主张合同系伪造",
                "priority": "核心证据",
                "tag": "推断",
                "q5_reinforce": "申请笔迹鉴定",
                "q6_failure_impact": "借贷关系无法成立",
            },
            {
                "evidence_id": "ev-002",
                "q1_what": "微信聊天记录截图",
                "q2_target": "证明催款行为",
                "q3_key_risk": "真实性存疑",
                "q4_best_attack": "主张截图不完整",
                "priority": "辅助证据",
                "tag": "推断",
            },
        ],
        "unified_electronic_strategy": "对所有电子证据申请公证",
        "evidence_battle_matrix": [],
        "scenario_tree": None,
    },
    "layer3": {
        "outputs": [
            {
                "perspective": "plaintiff",
                "evidence_supplement_checklist": ["补充银行流水"],
                "cross_examination_points": ["质疑被告对账单"],
                "trial_questions": ["请问被告何时还款?"],
                "contingency_plans": ["如证据不足则和解"],
                "over_assertion_boundaries": ["不得主张精神损害"],
                "unified_electronic_evidence_strategy": "",
            }
        ]
    },
    "layer4": {
        "adversarial_transcripts_md": "## 第一轮\n原告: 借款事实清楚。",
        "evidence_index_md": "| 证据 | 说明 |\n|---|---|\n| ev-001 | 借款合同 |",
        "timeline_md": "",
        "glossary_md": "## 术语\n借贷关系: 民法典第667条",
        "amount_calculation_md": "## 金额\n本金: 100,000元",
    },
    "created_at": "2024-01-01T00:00:00Z",
}


# ---------------------------------------------------------------------------
# Test class
# ---------------------------------------------------------------------------


class TestDocxV3Render:
    def test_smoke(self, tmp_path: Path) -> None:
        """Minimal report dict generates a .docx file with non-zero size."""
        minimal: dict = {
            "report_id": "rpt-min",
            "case_id": "case-min",
            "run_id": "run-min",
            "perspective": "neutral",
            "layer1": {
                "cover_summary": {"neutral_conclusion": "测试结论"},
            },
            "layer2": {},
            "layer3": {"outputs": []},
            "layer4": {},
        }
        dest = generate_docx_v3_report(output_dir=tmp_path, report_v3=minimal)

        assert dest.exists(), "DOCX file was not created"
        assert dest.stat().st_size > 0, "DOCX file is empty"
        assert dest.suffix == ".docx", f"Expected .docx suffix, got {dest.suffix}"

    def test_layer1_sections_present(self, tmp_path: Path) -> None:
        """Layer 1 renders the main section heading and timeline sub-heading."""
        dest = generate_docx_v3_report(output_dir=tmp_path, report_v3=_FULL_REPORT)
        text = _read_docx_text(dest)

        assert "一、封面摘要" in text, "Missing '一、封面摘要' heading"
        assert "案件时间线" in text, "Missing '案件时间线' sub-heading in layer1"

    def test_layer1_winning_move(self, tmp_path: Path) -> None:
        """When winning_move is populated, '胜负手' heading appears."""
        dest = generate_docx_v3_report(output_dir=tmp_path, report_v3=_FULL_REPORT)
        text = _read_docx_text(dest)

        assert "胜负手" in text, "Missing '胜负手' heading for populated winning_move"

    def test_layer2_evidence_cards_core(self, tmp_path: Path) -> None:
        """Core 6-field evidence card (ev-001) renders with '2.4 证据卡片' heading
        and '5. 如何补强' row label inside the table.
        """
        dest = generate_docx_v3_report(output_dir=tmp_path, report_v3=_FULL_REPORT)
        text = _read_docx_text(dest)

        assert "2.4 证据卡片" in text, "Missing '2.4 证据卡片' heading"
        assert "5. 如何补强" in text, "Missing '5. 如何补强' table row for core card ev-001"

    def test_layer2_evidence_cards_basic(self, tmp_path: Path) -> None:
        """Supporting card (ev-002, no q5/q6) renders a 4-row table only."""
        dest = generate_docx_v3_report(output_dir=tmp_path, report_v3=_FULL_REPORT)
        doc = Document(dest)

        # Count tables that have exactly 4 rows (basic card format)
        four_row_tables = [t for t in doc.tables if len(t.rows) == 4]
        assert len(four_row_tables) >= 1, (
            "Expected at least one 4-row table for the basic evidence card (ev-002)"
        )

        # ev-002 content should appear in the document
        text = _read_docx_text(dest)
        assert "微信聊天记录截图" in text, "ev-002 q1_what text not found in DOCX"

    def test_layer2_unified_strategy(self, tmp_path: Path) -> None:
        """Unified electronic evidence strategy heading renders when populated."""
        dest = generate_docx_v3_report(output_dir=tmp_path, report_v3=_FULL_REPORT)
        text = _read_docx_text(dest)

        assert "2.3 统一电子证据补强策略" in text, (
            "Missing '2.3 统一电子证据补强策略' heading"
        )

    def test_layer4_glossary(self, tmp_path: Path) -> None:
        """Layer 4 glossary section heading appears when glossary_md is set."""
        dest = generate_docx_v3_report(output_dir=tmp_path, report_v3=_FULL_REPORT)
        text = _read_docx_text(dest)

        assert "4.4 术语表" in text, "Missing '4.4 术语表' heading in layer4"

    def test_layer4_amount_calculation(self, tmp_path: Path) -> None:
        """Layer 4 amount calculation section appears when amount_calculation_md is set."""
        dest = generate_docx_v3_report(output_dir=tmp_path, report_v3=_FULL_REPORT)
        text = _read_docx_text(dest)

        assert "4.5 金额计算明细" in text, "Missing '4.5 金额计算明细' heading in layer4"

    def test_custom_filename(self, tmp_path: Path) -> None:
        """The filename parameter controls the output file name."""
        custom_name = "my_custom_report.docx"
        dest = generate_docx_v3_report(
            output_dir=tmp_path,
            report_v3=_FULL_REPORT,
            filename=custom_name,
        )

        assert dest.name == custom_name, (
            f"Expected filename '{custom_name}', got '{dest.name}'"
        )
        assert dest.exists(), "Custom-named DOCX file was not created"

    def test_docx_lint_passes(self, tmp_path: Path) -> None:
        """Generated DOCX passes the render contract lint with no ERROR-level results."""
        dest = generate_docx_v3_report(output_dir=tmp_path, report_v3=_FULL_REPORT)

        results = lint_docx_render_contract(dest)
        errors = [r for r in results if r.severity == LintSeverity.ERROR]
        assert errors == [], f"DOCX lint errors: {errors}"
