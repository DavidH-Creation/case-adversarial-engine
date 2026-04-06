from __future__ import annotations

import re
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch

from docx import Document

from engines.report_generation.docx_generator import (
    generate_docx_report,
    generate_docx_v3_report,
)


_EMPTY_RESULT: dict = {
    "case_id": "case-smoke-001",
    "run_id": "run-smoke-001",
    "rounds": [],
    "evidence_conflicts": [],
    "summary": None,
    "missing_evidence_report": [],
}

_EMPTY_CASE_DATA: dict = {
    "parties": {},
    "summary": [],
    "model": "",
}


def _read_docx_text(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _count_docx_headings(path: Path) -> int:
    """Count all heading paragraphs in a DOCX document."""
    doc = Document(path)
    count = 0
    for para in doc.paragraphs:
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            count += 1
    return count


def _count_md_h2(md_text: str) -> int:
    """Count ## headings in Markdown text."""
    return len(re.findall(r"(?m)^##\s+", md_text))


class TestDocxGeneratorSmoke:
    def test_minimal_data_no_exception(self, tmp_path: Path) -> None:
        dest = generate_docx_report(
            output_dir=tmp_path,
            case_data=_EMPTY_CASE_DATA,
            result=_EMPTY_RESULT,
        )
        assert dest.exists()
        assert dest.stat().st_size > 0
        assert dest.suffix == ".docx"

    def test_output_in_correct_directory(self, tmp_path: Path) -> None:
        dest = generate_docx_report(
            output_dir=tmp_path,
            case_data=_EMPTY_CASE_DATA,
            result=_EMPTY_RESULT,
        )
        assert dest.parent == tmp_path

    def test_custom_filename(self, tmp_path: Path) -> None:
        dest = generate_docx_report(
            output_dir=tmp_path,
            case_data=_EMPTY_CASE_DATA,
            result=_EMPTY_RESULT,
            filename="custom_report.docx",
        )
        assert dest.name == "custom_report.docx"
        assert dest.exists()

    def test_full_data_no_exception(self, tmp_path: Path) -> None:
        case_data = {
            "parties": {
                "plaintiff": {"party_id": "party-p-001", "name": "Alice"},
                "defendant": {"party_id": "party-d-001", "name": "Bob"},
            },
            "summary": [
                ["Case Type", "Civil Loan Dispute"],
                ["Amount", "100,000"],
            ],
            "model": "claude-sonnet-4-6",
        }

        result = {
            "case_id": "case-smoke-002",
            "run_id": "run-smoke-002",
            "rounds": [
                {
                    "round_number": 1,
                    "phase": "claim",
                    "outputs": [
                        {
                            "agent_role_code": "plaintiff_agent",
                            "title": "Plaintiff position",
                            "body": "The defendant failed to repay the loan.",
                            "issue_ids": ["issue-001"],
                            "evidence_citations": ["ev-001"],
                            "risk_flags": [
                                {"flag_id": "own-weakness-001", "description": "Date is blurry"}
                            ],
                        },
                        {
                            "agent_role_code": "defendant_agent",
                            "title": "Defendant response",
                            "body": "The loan was already repaid.",
                            "issue_ids": ["issue-001"],
                            "evidence_citations": ["ev-002"],
                            "risk_flags": [],
                        },
                    ],
                }
            ],
            "evidence_conflicts": [
                {
                    "issue_id": "issue-001",
                    "conflict_description": "The parties dispute whether repayment happened.",
                }
            ],
            "summary": {
                "overall_assessment": "The plaintiff currently has the stronger documentary record.",
                "plaintiff_strongest_arguments": [
                    {
                        "issue_id": "issue-001",
                        "position": "The IOU is valid",
                        "reasoning": "The original is available",
                    }
                ],
                "defendant_strongest_defenses": [
                    {
                        "issue_id": "issue-001",
                        "position": "The debt was repaid",
                        "reasoning": "There is a transfer record",
                    }
                ],
                "unresolved_issues": [
                    {
                        "issue_id": "issue-001",
                        "issue_title": "Repayment dispute",
                        "why_unresolved": "The transfer note is ambiguous",
                    }
                ],
            },
            "missing_evidence_report": [
                {
                    "issue_id": "issue-001",
                    "missing_for_party_id": "party-d-001",
                    "description": "Repayment proof is incomplete",
                }
            ],
        }

        exec_summary = {
            "top5_decisive_issues": ["issue-001"],
            "current_most_stable_claim": "Repayment of principal in the amount of 100,000",
            "top3_immediate_actions": ["Authenticate the IOU", "Request bank records"],
            "critical_evidence_gaps": ["Repayment proof"],
            "top3_adversary_optimal_attacks": ["Challenge the IOU authenticity"],
        }

        issue = SimpleNamespace(
            issue_id="issue-001",
            title="Repayment dispute",
            issue_type=SimpleNamespace(value="factual"),
        )
        issue_tree = SimpleNamespace(issues=[issue])

        dest = generate_docx_report(
            output_dir=tmp_path,
            case_data=case_data,
            result=result,
            issue_tree=issue_tree,
            exec_summary=exec_summary,
            filename="full_report.docx",
        )

        assert dest.exists()
        assert dest.stat().st_size > 0


class TestDocxGeneratorProbabilityFree:
    def test_decision_tree_uses_path_ranking_without_probability_labels(
        self, tmp_path: Path
    ) -> None:
        dest = generate_docx_report(
            output_dir=tmp_path,
            case_data=_EMPTY_CASE_DATA,
            result=_EMPTY_RESULT,
            decision_tree={
                "paths": [
                    {
                        "path_id": "path-001",
                        "trigger_condition": "Trigger one",
                        "trigger_issue_ids": ["ISS-001"],
                        "key_evidence_ids": ["EV-001"],
                        "possible_outcome": "Outcome one",
                        "probability": 0.95,
                        "probability_rationale": "legacy only",
                        "party_favored": "plaintiff",
                    },
                    {
                        "path_id": "path-002",
                        "trigger_condition": "Trigger two",
                        "trigger_issue_ids": ["ISS-002"],
                        "key_evidence_ids": ["EV-002"],
                        "possible_outcome": "Outcome two",
                        "probability": 0.10,
                        "probability_rationale": "legacy only",
                        "party_favored": "defendant",
                    },
                ],
                "most_likely_path": "path-001",
                "plaintiff_best_path": "path-001",
                "defendant_best_path": "path-002",
                "path_ranking": [
                    {"path_id": "path-002", "probability": 0.10, "party_favored": "defendant"},
                    {"path_id": "path-001", "probability": 0.95, "party_favored": "plaintiff"},
                ],
                "blocking_conditions": [],
            },
            amount_report={
                "claim_calculation_table": [
                    {
                        "claim_id": "CLM-001",
                        "claimed_amount": "100000",
                        "calculated_amount": "80000",
                    }
                ],
                "consistency_check_result": {
                    "verdict_block_active": False,
                    "unresolved_conflicts": [],
                },
            },
        )

        content = _read_docx_text(dest)

        assert "调解区间评估" not in content
        assert "路径概率比较" not in content
        assert "可能性：" not in content
        assert "概率依据" not in content
        assert content.index("Outcome two") < content.index("Outcome one")


# ---------------------------------------------------------------------------
# V3 DOCX fixtures
# ---------------------------------------------------------------------------

_V3_REPORT_MINIMAL: dict = {
    "report_id": "rpt-v3-test001",
    "case_id": "case-test-v3",
    "run_id": "run-test-v3",
    "perspective": "neutral",
    "layer1": {
        "cover_summary": {
            "neutral_conclusion": "本案涉及3个争点。",
            "winning_move": "借条原件是本案胜负手。",
            "blocking_conditions": ["如被告提供完整还款凭证则结论逆转"],
        },
        "timeline": [
            {"date": "2025-01-01", "event": "签署借条", "source": "ev-001", "disputed": False},
        ],
        "evidence_priorities": [
            {"evidence_id": "ev-001", "title": "借条原件", "priority": "核心", "reason": "直接证明"},
        ],
    },
    "layer2": {
        "fact_base": [
            {"description": "双方确认借款事实", "source_evidence_ids": ["ev-001"]},
        ],
        "issue_map": [
            {
                "issue_id": "issue-001",
                "issue_title": "借款合同效力",
                "depth": 0,
                "plaintiff_thesis": "合同有效",
                "defendant_thesis": "合同无效",
                "decisive_evidence": ["ev-001"],
                "current_gaps": [],
                "outcome_sensitivity": "高",
            },
        ],
        "evidence_cards": [
            {
                "evidence_id": "ev-001",
                "q1_what": "借条",
                "q2_target": "证明借款关系",
                "q3_key_risk": "笔迹鉴定",
                "q4_best_attack": "质疑签名真实性",
                "q5_reinforce": "公证",
                "q6_failure_impact": "丧失核心证据",
                "priority": "核心",
            },
        ],
    },
    "layer3": {
        "outputs": [
            {
                "perspective": "plaintiff",
                "evidence_supplement_checklist": ["补充银行流水"],
                "cross_examination_points": ["质证借条签名"],
                "trial_questions": ["询问还款时间"],
                "contingency_plans": ["备选：请求笔迹鉴定"],
                "over_assertion_boundaries": ["不宜主张精神损害赔偿"],
            },
            {
                "perspective": "defendant",
                "evidence_supplement_checklist": ["提供还款转账记录"],
                "cross_examination_points": ["质证借条日期"],
                "trial_questions": ["询问借款用途"],
                "contingency_plans": ["备选：和解方案"],
                "over_assertion_boundaries": ["不宜否认全部借款事实"],
            },
        ],
    },
    "layer4": {
        "adversarial_transcripts_md": "## 第一轮\n\n原告主张...\n\n被告抗辩...",
        "evidence_index_md": "| 证据 | 来源 |\n|---|---|\n| 借条 | 原告 |",
        "glossary_md": "- **借条**: 债权凭证",
    },
}


# ---------------------------------------------------------------------------
# Task 1: DOCX content completeness — section count parity with MD
# ---------------------------------------------------------------------------


class TestDocxV3ContentCompleteness:
    """Assert V3 DOCX heading count >= MD ## heading count."""

    def test_v3_docx_section_count_gte_md(self, tmp_path: Path) -> None:
        """DOCX must contain at least as many heading sections as MD ## count."""
        dest = generate_docx_v3_report(
            output_dir=tmp_path,
            report_v3=_V3_REPORT_MINIMAL,
        )
        assert dest.exists()

        docx_headings = _count_docx_headings(dest)

        # The MD report for the same V3 data would have these ## sections:
        # A, B, C, D, E (layer1) + 2.1, 2.2, 2.4 (layer2) +
        # plaintiff策略, defendant策略 (layer3) + 4.1, 4.2, 4.4 (layer4)
        # That's about 12-13 ## headings. DOCX may have more (level-3 sub-headings).
        # Core assertion: DOCX heading count >= expected MD ## count.
        expected_md_h2_count = 12  # conservative lower bound
        assert docx_headings >= expected_md_h2_count, (
            f"DOCX has {docx_headings} headings but expected >= {expected_md_h2_count}"
        )

    def test_v3_docx_contains_all_layer_headings(self, tmp_path: Path) -> None:
        """DOCX must contain headings for all four layers."""
        dest = generate_docx_v3_report(
            output_dir=tmp_path,
            report_v3=_V3_REPORT_MINIMAL,
        )
        content = _read_docx_text(dest)

        layer_markers = [
            "一、封面摘要",
            "二、中立对抗内核",
            "三、角色化输出",
            "四、附录",
        ]
        for marker in layer_markers:
            assert marker in content, f"Missing layer heading: {marker}"

    def test_v3_docx_contains_layer1_subsections(self, tmp_path: Path) -> None:
        """DOCX layer1 should render A/B/C/D/E subsections when data is present."""
        dest = generate_docx_v3_report(
            output_dir=tmp_path,
            report_v3=_V3_REPORT_MINIMAL,
        )
        content = _read_docx_text(dest)

        subsections = [
            "中立结论摘要",
            "胜负手",
            "阻断条件",
            "案件时间线",
            "证据优先级",
        ]
        for sub in subsections:
            assert sub in content, f"Missing layer1 subsection: {sub}"

    def test_v3_docx_contains_layer2_subsections(self, tmp_path: Path) -> None:
        """DOCX layer2 should render 2.1 fact_base, 2.2 issue_map, 2.4 evidence_cards."""
        dest = generate_docx_v3_report(
            output_dir=tmp_path,
            report_v3=_V3_REPORT_MINIMAL,
        )
        content = _read_docx_text(dest)

        assert "事实底座" in content
        assert "争点地图" in content
        assert "证据卡片" in content

    def test_v3_docx_contains_layer3_perspectives(self, tmp_path: Path) -> None:
        """DOCX layer3 should have both plaintiff and defendant strategy sections."""
        dest = generate_docx_v3_report(
            output_dir=tmp_path,
            report_v3=_V3_REPORT_MINIMAL,
        )
        content = _read_docx_text(dest)

        assert "原告策略" in content
        assert "被告策略" in content

    def test_v3_docx_contains_layer4_appendix_items(self, tmp_path: Path) -> None:
        """DOCX layer4 should render transcripts, evidence index, glossary."""
        dest = generate_docx_v3_report(
            output_dir=tmp_path,
            report_v3=_V3_REPORT_MINIMAL,
        )
        content = _read_docx_text(dest)

        assert "三轮对抗辩论记录" in content
        assert "证据索引" in content
        assert "术语表" in content


# ---------------------------------------------------------------------------
# Task 2: V3.0 backward-compat fallback coverage
# ---------------------------------------------------------------------------


class TestDocxV3FallbackPaths:
    """Verify V3.0 fallback branches still render correctly for legacy data."""

    def test_v30_plaintiff_defendant_summary_fallback(self, tmp_path: Path) -> None:
        """When winning_move is absent, V3.0 plaintiff/defendant summaries render."""
        report = {
            **_V3_REPORT_MINIMAL,
            "layer1": {
                "cover_summary": {
                    "neutral_conclusion": "测试结论。",
                    "winning_move": "",  # empty → triggers V3.0 fallback
                    "plaintiff_summary": {
                        "top3_strengths": ["优势一", "优势二", "优势三"],
                        "top2_dangers": ["风险一", "风险二"],
                    },
                    "defendant_summary": {
                        "top3_defenses": ["防线一", "防线二", "防线三"],
                        "plaintiff_likely_supplement": [],
                        "optimal_attack_order": [],
                    },
                },
            },
            "layer2": {"fact_base": [], "issue_map": []},
            "layer3": {"outputs": []},
            "layer4": {},
        }
        dest = generate_docx_v3_report(output_dir=tmp_path, report_v3=report)
        content = _read_docx_text(dest)

        assert "原告视角" in content
        assert "被告视角" in content
        assert "优势1" in content or "优势一" in content

    def test_v30_evidence_traffic_lights_fallback(self, tmp_path: Path) -> None:
        """When evidence_priorities is absent, V3.0 traffic lights render."""
        report = {
            **_V3_REPORT_MINIMAL,
            "layer1": {
                "cover_summary": {
                    "neutral_conclusion": "结论。",
                    "winning_move": "胜负手",
                    "blocking_conditions": [],
                },
                "evidence_priorities": [],  # empty → triggers V3.0 fallback
                "evidence_traffic_lights": [
                    {
                        "evidence_id": "ev-001",
                        "title": "借条",
                        "risk_level": "green",
                        "reason": "原件完整",
                    },
                ],
            },
            "layer2": {"fact_base": [], "issue_map": []},
            "layer3": {"outputs": []},
            "layer4": {},
        }
        dest = generate_docx_v3_report(output_dir=tmp_path, report_v3=report)
        content = _read_docx_text(dest)

        assert "证据风险红绿灯" in content
        assert "借条" in content

    def test_v30_evidence_battle_matrix_fallback(self, tmp_path: Path) -> None:
        """When evidence_cards is absent, V3.0 evidence_battle_matrix renders."""
        report = {
            **_V3_REPORT_MINIMAL,
            "layer1": {
                "cover_summary": {"neutral_conclusion": "结论。"},
            },
            "layer2": {
                "fact_base": [],
                "issue_map": [],
                "evidence_cards": [],  # empty → triggers V3.0 fallback
                "evidence_battle_matrix": [
                    {
                        "evidence_id": "ev-001",
                        "risk_level": "yellow",
                        "q1_what": "借条原件",
                        "q2_proves": "借款事实",
                        "q3_direction": "有利原告",
                        "q4_risks": "笔迹质疑",
                        "q5_opponent_attack": "否认签名",
                        "q6_reinforce": "公证",
                        "q7_failure_impact": "核心证据缺失",
                    },
                ],
            },
            "layer3": {"outputs": []},
            "layer4": {},
        }
        dest = generate_docx_v3_report(output_dir=tmp_path, report_v3=report)
        content = _read_docx_text(dest)

        assert "证据作战矩阵" in content
        assert "借条原件" in content


# ---------------------------------------------------------------------------
# Task 3: CJK font fallback — no crash without Microsoft YaHei
# ---------------------------------------------------------------------------


class TestDocxCJKFontFallback:
    """DOCX generation must not crash when CJK font is unavailable."""

    def test_generate_without_microsoft_yahei(self, tmp_path: Path) -> None:
        """Smoke test: DOCX generation succeeds even if YaHei is missing.

        python-docx embeds the font name as metadata only; it does not validate
        font availability at generation time.  This test ensures the generator
        doesn't do any font-existence checks that would raise.
        """
        dest = generate_docx_v3_report(
            output_dir=tmp_path,
            report_v3=_V3_REPORT_MINIMAL,
        )
        assert dest.exists()
        assert dest.stat().st_size > 0

        # Verify the DOCX can be reopened and has CJK content
        doc = Document(dest)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "封面摘要" in full_text

    def test_legacy_docx_without_microsoft_yahei(self, tmp_path: Path) -> None:
        """Legacy generator also works without YaHei font."""
        dest = generate_docx_report(
            output_dir=tmp_path,
            case_data=_EMPTY_CASE_DATA,
            result=_EMPTY_RESULT,
        )
        assert dest.exists()
        assert dest.stat().st_size > 0

    def test_font_fallback_chain_in_docx_xml(self, tmp_path: Path) -> None:
        """DOCX rFonts should specify eastAsia font (even if unavailable at runtime)."""
        dest = generate_docx_v3_report(
            output_dir=tmp_path,
            report_v3=_V3_REPORT_MINIMAL,
        )
        doc = Document(dest)
        # Check that at least one run has eastAsia font set
        found_east_asia = False
        for para in doc.paragraphs:
            for run in para.runs:
                rPr = run._element.find(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr"
                )
                if rPr is not None:
                    rFonts = rPr.find(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts"
                    )
                    if rFonts is not None:
                        ea = rFonts.get(
                            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia"
                        )
                        if ea:
                            found_east_asia = True
                            break
            if found_east_asia:
                break
        assert found_east_asia, "No eastAsia font attribute found in DOCX runs"


# ---------------------------------------------------------------------------
# Task 4: V3.0 Layer3 role-specific fallback
# ---------------------------------------------------------------------------


class TestDocxV3Layer3Fallback:
    """Verify V3.0 role-specific rendering when V3.1 action fields are absent."""

    def test_v30_plaintiff_role_fallback(self, tmp_path: Path) -> None:
        """V3.0 plaintiff rendering: top_claims, evidence_to_supplement, etc."""
        report = {
            **_V3_REPORT_MINIMAL,
            "layer1": {"cover_summary": {"neutral_conclusion": "测试。"}},
            "layer2": {"fact_base": [], "issue_map": []},
            "layer3": {
                "outputs": [
                    {
                        "perspective": "plaintiff",
                        # No V3.1 action fields → triggers V3.0 fallback
                        "top_claims": ["诉请一", "诉请二"],
                        "defendant_attack_chains": ["攻击链预警"],
                        "evidence_to_supplement": ["需补强证据"],
                        "trial_sequence": ["第一步举证"],
                        "claims_to_abandon": ["放弃诉请"],
                    },
                ],
            },
            "layer4": {},
        }
        dest = generate_docx_v3_report(output_dir=tmp_path, report_v3=report)
        content = _read_docx_text(dest)

        assert "三大诉请" in content
        assert "诉请一" in content
        assert "被告攻击链预警" in content
        assert "需补强证据清单" in content
        assert "庭审举证顺序" in content
        assert "应放弃诉请" in content

    def test_v30_defendant_role_fallback(self, tmp_path: Path) -> None:
        """V3.0 defendant rendering: top_defenses, motions_to_file, etc."""
        report = {
            **_V3_REPORT_MINIMAL,
            "layer1": {"cover_summary": {"neutral_conclusion": "测试。"}},
            "layer2": {"fact_base": [], "issue_map": []},
            "layer3": {
                "outputs": [
                    {
                        "perspective": "defendant",
                        # No V3.1 action fields → triggers V3.0 fallback
                        "top_defenses": ["防线一", "防线二"],
                        "plaintiff_supplement_prediction": ["补强预测"],
                        "evidence_to_challenge_first": ["优先质证"],
                        "motions_to_file": ["申请鉴定"],
                        "over_assertion_warnings": ["过度主张"],
                    },
                ],
            },
            "layer4": {},
        }
        dest = generate_docx_v3_report(output_dir=tmp_path, report_v3=report)
        content = _read_docx_text(dest)

        assert "三大防线" in content
        assert "防线一" in content
        assert "原告可能补强方向" in content
        assert "优先质证目标" in content
        assert "应提交动议" in content
        assert "过度主张警告" in content


# ---------------------------------------------------------------------------
# Task 5: Layer4 timeline placeholder skip
# ---------------------------------------------------------------------------


class TestDocxV3Layer4TimelinePlaceholder:
    """Verify that Layer4 timeline_md is skipped when it only contains placeholder text."""

    def test_timeline_placeholder_is_skipped(self, tmp_path: Path) -> None:
        """timeline_md with '暂无时间线数据' should not render in Layer4."""
        report = {
            **_V3_REPORT_MINIMAL,
            "layer4": {
                "timeline_md": "暂无时间线数据",
                "glossary_md": "- **测试**: 术语",
            },
        }
        dest = generate_docx_v3_report(output_dir=tmp_path, report_v3=report)
        content = _read_docx_text(dest)

        # Timeline heading should NOT appear
        assert "4.3 案件时间线" not in content
        # But glossary should still render
        assert "术语表" in content

    def test_real_timeline_does_render(self, tmp_path: Path) -> None:
        """timeline_md with real content should render normally."""
        report = {
            **_V3_REPORT_MINIMAL,
            "layer4": {
                "timeline_md": "## 时间线\n\n2025-01-01 签署借条",
                "glossary_md": "- **测试**: 术语",
            },
        }
        dest = generate_docx_v3_report(output_dir=tmp_path, report_v3=report)
        content = _read_docx_text(dest)

        assert "案件时间线" in content
        assert "签署借条" in content


# ---------------------------------------------------------------------------
# Task 6: V3 empty data resilience
# ---------------------------------------------------------------------------


class TestDocxV3EmptyDataResilience:
    """V3 DOCX generation must not crash with entirely empty layer data."""

    def test_all_layers_empty(self, tmp_path: Path) -> None:
        """Completely empty report_v3 must produce a valid DOCX without error."""
        report = {
            "case_id": "case-empty",
            "run_id": "run-empty",
            "perspective": "neutral",
            "layer1": {},
            "layer2": {},
            "layer3": {},
            "layer4": {},
        }
        dest = generate_docx_v3_report(output_dir=tmp_path, report_v3=report)
        assert dest.exists()
        assert dest.stat().st_size > 0

        # Must still have the 4-layer headings
        content = _read_docx_text(dest)
        assert "一、封面摘要" in content
        assert "二、中立对抗内核" in content
        assert "三、角色化输出" in content
        assert "四、附录" in content

    def test_layer2_empty_fact_base_shows_fallback_text(self, tmp_path: Path) -> None:
        """Empty fact_base should show '暂无' fallback text, not crash."""
        report = {
            **_V3_REPORT_MINIMAL,
            "layer2": {"fact_base": [], "issue_map": [], "evidence_cards": []},
        }
        dest = generate_docx_v3_report(output_dir=tmp_path, report_v3=report)
        content = _read_docx_text(dest)
        assert "暂无双方均认可的无争议事实" in content


# ---------------------------------------------------------------------------
# Task 7: V3.0 scenario_tree_summary fallback (Layer1 C section)
# ---------------------------------------------------------------------------


class TestDocxV3ScenarioTreeFallback:
    """Verify V3.0 scenario_tree_summary fallback when blocking_conditions is empty."""

    def test_scenario_tree_fallback_renders(self, tmp_path: Path) -> None:
        """When blocking_conditions is empty, scenario_tree_summary should render."""
        report = {
            **_V3_REPORT_MINIMAL,
            "layer1": {
                "cover_summary": {
                    "neutral_conclusion": "结论",
                    "winning_move": "胜负手",
                    "blocking_conditions": [],  # empty → V3.0 fallback
                },
                "scenario_tree_summary": "条件一；条件二；条件三",
            },
            "layer2": {"fact_base": [], "issue_map": []},
            "layer3": {"outputs": []},
            "layer4": {},
        }
        dest = generate_docx_v3_report(output_dir=tmp_path, report_v3=report)
        content = _read_docx_text(dest)

        assert "条件场景摘要" in content
        assert "条件一" in content
        assert "条件二" in content
