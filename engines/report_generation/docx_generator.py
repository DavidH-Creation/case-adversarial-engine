"""
通用 Word 文档报告生成器 — 对抗分析报告模板。
Reusable DOCX report generator for adversarial case analysis.

从 run_case.py 管道末端调用，接收结构化数据，生成统一格式的 Word 报告。
Zero hardcoded case-specific content — all data comes from parameters.

Usage:
    from engines.report_generation.docx_generator import generate_docx_report
    dest = generate_docx_report(
        output_dir=Path("outputs/20260329-123456"),
        case_data=case_data,          # YAML dict
        result=result_dict,           # AdversarialResult as dict
        issue_tree=issue_tree,        # IssueTree object (has .issues with .title, .issue_type)
        decision_tree=dt_dict,        # DecisionPathTree as dict (or None)
        attack_chain=ac_dict,         # OptimalAttackChain as dict (or None)
        exec_summary=es_dict,         # ExecutiveSummaryArtifact as dict (or None)
        amount_report=ar_dict,        # AmountCalculationReport as dict (or None)
    )
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Emu

from engines.report_generation.mediation_range import compute_mediation_range
from engines.shared.disclaimer_templates import DISCLAIMER_DOCX_BODY, DISCLAIMER_DOCX_TITLE
from engines.report_generation.risk_heatmap import build_risk_heatmap, RISK_LABEL_ZH, RiskLevel
from engines.report_generation.v3.tag_system import humanize_text
from engines.report_generation.v3.docx_styles import (
    CLR_TITLE_DARK, CLR_BLUE, CLR_RED, CLR_GREEN, CLR_ORANGE, CLR_BODY, CLR_GRAY,
    SZ_TITLE, SZ_SUBTITLE, SZ_AGENT_TITLE, SZ_SECTION_HDR, SZ_BODY, SZ_RISK, SZ_EVIDENCE, SZ_NORMAL,
    FONT_NAME, FONT_EAST_ASIA,
)

# ---------------------------------------------------------------------------
# ID humanization context (populated per-report)
# ---------------------------------------------------------------------------
_humanize_ctx: dict[str, str] = {}


def _h(text) -> str:
    """Humanize internal IDs in text for user-facing output."""
    if not text:
        return ""
    return humanize_text(str(text), context=_humanize_ctx)


# ---------------------------------------------------------------------------
# 翻译映射 / Translation maps
# ---------------------------------------------------------------------------
_ISSUE_TYPE_ZH: dict[str, str] = {
    "factual": "事实争点",
    "mixed": "混合争点",
    "legal": "法律争点",
    "procedural": "程序争点",
}

_ROUND_PHASE_ZH: dict[str, str] = {
    "claim": "首轮主张",
    "evidence": "证据整理",
    "rebuttal": "针对性反驳",
}


# ---------------------------------------------------------------------------
# 内部工具函数 / Internal helpers
# ---------------------------------------------------------------------------


def _add_run(
    para,
    text: str,
    *,
    bold: bool = False,
    size: int | None = None,
    color: RGBColor | None = None,
    italic: bool = False,
):
    """向段落追加一个格式化 run。"""
    text = _h(text)
    run = para.add_run(text)
    run.font.name = FONT_NAME
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    if bold:
        run.font.bold = True
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color
    if italic:
        run.font.italic = True
    return run


def _styled(
    doc, text: str, *, bold: bool = False, size: int = SZ_NORMAL, color: RGBColor = CLR_BODY
):
    """添加单 run 段落。"""
    p = doc.add_paragraph()
    _add_run(p, text, bold=bold, size=size, color=color)
    return p


def _bullet(doc, text: str, *, size: int = SZ_NORMAL, color: RGBColor = CLR_BODY):
    p = doc.add_paragraph()
    _add_run(p, "• " + text, size=size, color=color)
    return p


def _agent_color(role: str) -> RGBColor:
    if "plaintiff" in role:
        return CLR_BLUE
    if "defendant" in role:
        return CLR_RED
    return CLR_GREEN


def _set_table_font(table):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = FONT_NAME
                    run.font.size = SZ_NORMAL


def _build_party_zh(case_data: dict) -> dict[str, str]:
    """从 YAML case_data 动态构建 party_id → 中文名映射。"""
    mapping: dict[str, str] = {}
    parties = case_data.get("parties", {})
    for role_key, info in parties.items():
        pid = info.get("party_id", "")
        name = info.get("name", "")
        if not pid or not name:
            continue
        if "plaintiff" in role_key:
            mapping[pid] = f"原告{name}方"
        else:
            mapping[pid] = f"被告{name}方"
    return mapping


def _build_issue_info(issue_tree) -> dict[str, tuple[str, str]]:
    """从 IssueTree 对象构建 issue_id → (title, type_value) 映射。"""
    info: dict[str, tuple[str, str]] = {}
    if issue_tree is None:
        return info
    for iss in getattr(issue_tree, "issues", []):
        itype = (
            getattr(iss.issue_type, "value", str(iss.issue_type))
            if hasattr(iss, "issue_type")
            else ""
        )
        info[iss.issue_id] = (iss.title, itype)
    return info


_UUID_RE = re.compile(
    r"^[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}$",
    re.IGNORECASE,
)


def _is_uuid(s: str) -> bool:
    """Return True if s looks like a raw UUID string."""
    return bool(_UUID_RE.match(s))


def _filter_uuids(items: list[str]) -> list[str]:
    """Remove raw UUID strings from a list of text items."""
    return [item for item in items if not _is_uuid(item)]


# ---------------------------------------------------------------------------
# 公共 API / Public API
# ---------------------------------------------------------------------------


def generate_docx_report(
    *,
    output_dir: Path,
    case_data: dict,
    result: dict,
    issue_tree: Any = None,
    ranked_issues: Any = None,
    decision_tree: dict | None = None,
    attack_chain: dict | None = None,
    exec_summary: dict | None = None,
    amount_report: dict | None = None,
    action_rec: Any = None,
    document_drafts: list | None = None,
    similar_cases: list | None = None,
    filename: str | None = None,
) -> Path:
    """生成通用对抗分析 Word 报告。

    Args:
        output_dir:      输出目录
        case_data:       YAML 案件定义 dict（含 parties, summary, case_type 等）
        result:          AdversarialResult 序列化 dict
        issue_tree:      IssueTree 对象（可选，用于争点描述/类型）
        decision_tree:   DecisionPathTree 序列化 dict（可选）
        attack_chain:    OptimalAttackChain 序列化 dict（可选）
        exec_summary:    ExecutiveSummaryArtifact 序列化 dict（可选）
        amount_report:   AmountCalculationReport 序列化 dict（可选）
        similar_cases:   RankedCase 序列化 dict 列表（可选，类案检索结果）
        filename:        自定义文件名（默认 "对抗分析报告.docx"）

    Returns:
        生成的 docx 文件路径
    """
    decision_tree = decision_tree or {}
    attack_chain = attack_chain or {}
    exec_summary = exec_summary or {}
    amount_report = amount_report or {}

    global _humanize_ctx
    _humanize_ctx = {}
    # Build humanize context from issue_tree
    if issue_tree:
        issues = getattr(issue_tree, "issues", [])
        for iss in issues:
            iid = getattr(iss, "issue_id", "")
            title = getattr(iss, "title", "")
            if iid and title:
                _humanize_ctx[iid] = title
    # Build from evidence in result
    for rnd in (result or {}).get("rounds", []):
        for ev in rnd.get("evidence_index", []):
            eid = ev.get("evidence_id", "")
            etitle = ev.get("title", "")
            if eid and etitle:
                _humanize_ctx[eid] = etitle

    party_zh = _build_party_zh(case_data)
    issue_info = _build_issue_info(issue_tree)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Emu(SZ_NORMAL)

    # ── 标题块 ──
    _render_title(doc, case_data, result)
    # ── 免责声明 ──
    _render_disclaimer(doc)
    # ── 行动优先级清单 ──
    _render_action_priority_list(doc, exec_summary, action_rec)
    # ── 案件摘要 ──
    _render_case_summary(doc, case_data, result)
    # ── 争点列表 ──
    _render_issue_table(doc, result, issue_info)
    # ── 三轮对抗 ──
    _render_debate_rounds(doc, result)
    # ── 证据冲突 ──
    _render_conflicts(doc, result)
    # ── AI 综合分析 ──
    _render_llm_summary(doc, result)
    # ── 证据缺失 ──
    _render_missing_evidence(doc, result, party_zh)
    # ── 争点影响排序 ──
    _render_issue_ranking(doc, exec_summary)
    # ── 风险热力图 ──
    _render_risk_heatmap(doc, ranked_issues)
    # ── 裁判路径树 ──
    _render_decision_tree_probability_free(doc, decision_tree)
    # ── 调解区间 ──
    # Mainline output no longer renders mediation sections.
    # ── 攻击链 ── (removed: now unified in _render_opponent_strategy_warning)
    # ── 对方策略预警 ──
    _render_opponent_strategy_warning(doc, result, attack_chain)
    # ── 行动建议 ──
    _render_action_recommendations(doc, exec_summary)
    # ── 执行摘要 ──
    _render_executive_summary(doc, exec_summary, amount_report)
    # ── 文书草稿 ──
    if document_drafts:
        _render_document_drafts(doc, document_drafts)
    # ── 类案检索参考 ──
    if similar_cases:
        _render_similar_cases(doc, similar_cases)

    # 保存
    if filename is None:
        filename = "对抗分析报告.docx"
    dest = output_dir / filename
    doc.save(str(dest))
    return dest


# ---------------------------------------------------------------------------
# 各章节渲染函数 / Section renderers
# ---------------------------------------------------------------------------


def _render_disclaimer(doc):
    """免责声明块（首页标题后）。"""
    _styled(doc, DISCLAIMER_DOCX_TITLE, bold=True, size=SZ_SECTION_HDR, color=CLR_ORANGE)
    _styled(doc, DISCLAIMER_DOCX_BODY, size=SZ_NORMAL, color=CLR_GRAY)
    doc.add_paragraph()


def _render_title(doc, case_data: dict, result: dict):
    """封面标题块。"""
    case_id = result.get("case_id", "")
    # 从 case_id 提取简短标识，如 "case-civil-loan-wang-v-chen-zhuang-2025" → "对抗分析报告"
    p0 = doc.add_paragraph()
    _add_run(p0, "对抗分析报告", bold=True, size=SZ_TITLE, color=CLR_TITLE_DARK)

    # 副标题：从 summary 提取核心争议 或用当事人名
    parties = case_data.get("parties", {})
    p_name = parties.get("plaintiff", {}).get("name", "原告")
    d_name = parties.get("defendant", {}).get("name", "被告")
    subtitle = f"{p_name} 诉 {d_name}"
    # 如果 summary 里有核心争议，追加
    for row in case_data.get("summary", []):
        if isinstance(row, list) and len(row) >= 2 and "争议" in str(row[0]):
            subtitle = f"{row[1]}（{subtitle}）"
            break

    p1 = doc.add_paragraph()
    _add_run(p1, subtitle, size=SZ_SUBTITLE, color=CLR_BLUE)

    p2 = doc.add_paragraph()
    model = case_data.get("model", "")
    _add_run(p2, f"对抗分析报告  |  对抗引擎  |  {model}", size=SZ_NORMAL, color=CLR_GRAY)
    doc.add_paragraph()


def _render_case_summary(doc, case_data: dict, result: dict):
    """案件摘要表格。"""
    doc.add_heading("案件摘要", level=1)

    rows = [
        ("案件ID", result.get("case_id", "")),
        ("运行ID", result.get("run_id", "")),
    ]
    # 追加 YAML summary 表
    for item in case_data.get("summary", []):
        if isinstance(item, list) and len(item) >= 2:
            rows.append((str(item[0]), str(item[1])))

    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(rows):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    _set_table_font(table)


def _render_issue_table(doc, result: dict, issue_info: dict[str, tuple[str, str]]):
    """争点列表表格。"""
    # 收集所有唯一 issue_id
    seen_ids: list[str] = []
    seen_set: set[str] = set()
    for r in result.get("rounds", []):
        for o in r.get("outputs", []):
            for iid in o.get("issue_ids", []):
                if iid not in seen_set:
                    seen_ids.append(iid)
                    seen_set.add(iid)

    conflict_ids = {c["issue_id"] for c in result.get("evidence_conflicts", [])}

    doc.add_heading(f"争点列表（{len(seen_ids)}个争点）", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "争点"
    hdr[2].text = "类型"
    hdr[3].text = "状态"

    for idx, iid in enumerate(seen_ids, 1):
        row = table.add_row().cells
        num = iid.split("-")[-1] if "-" in iid else str(idx)
        row[0].text = num
        desc, itype = issue_info.get(iid, (_h(iid), ""))
        row[1].text = desc
        row[2].text = _ISSUE_TYPE_ZH.get(itype, itype)
        row[3].text = "⚠ 冲突" if iid in conflict_ids else "✔"
    _set_table_font(table)


def _render_debate_rounds(doc, result: dict):
    """三轮对抗记录。"""
    doc.add_heading("三轮对抗记录", level=1)

    for r in result.get("rounds", []):
        rn = r["round_number"]
        phase = r["phase"]
        label = _ROUND_PHASE_ZH.get(phase, phase)
        doc.add_heading(f"第{rn}轮：{label}", level=2)

        for o in r["outputs"]:
            role = o["agent_role_code"]
            title_text = o.get("title", "")
            body = o.get("body", "")
            ev_cited = ", ".join(_h(x) for x in o.get("evidence_citations", []))

            if "plaintiff" in role:
                role_label = "[原告代理]"
            elif "defendant" in role:
                role_label = "[被告代理]"
            else:
                role_label = "[证据管理]"

            p = doc.add_paragraph()
            _add_run(
                p,
                f"{role_label} {title_text}",
                bold=True,
                size=SZ_AGENT_TITLE,
                color=_agent_color(role),
            )
            _styled(
                doc, body[:2000] + ("..." if len(body) > 2000 else ""), size=SZ_BODY, color=CLR_BODY
            )
            _styled(doc, f"引用证据: {ev_cited}", size=SZ_EVIDENCE, color=CLR_GRAY)

            risk_flags = o.get("risk_flags", [])
            if risk_flags:
                _styled(doc, "风险标记:", bold=True, size=SZ_RISK, color=CLR_ORANGE)
                for rf in risk_flags:
                    desc = rf.get("description", "")
                    flag_id = rf.get("flag_id", "")
                    prefix = "【本方薄弱点】" if flag_id.startswith("own-weakness-") else ""
                    _bullet(doc, prefix + desc, size=SZ_RISK, color=CLR_GRAY)

            doc.add_paragraph()


def _render_conflicts(doc, result: dict):
    """证据冲突分析。"""
    conflicts = result.get("evidence_conflicts", [])
    if not conflicts:
        return
    doc.add_heading(f"证据冲突分析（{len(conflicts)}条）", level=1)
    for c in conflicts:
        p = doc.add_paragraph()
        _add_run(p, f"[{_h(c['issue_id'])}] ", bold=True, size=SZ_SECTION_HDR, color=CLR_RED)
        doc.add_paragraph()
        _styled(doc, c["conflict_description"], size=SZ_NORMAL, color=CLR_BODY)


def _render_llm_summary(doc, result: dict):
    """AI 综合分析。"""
    summary = result.get("summary")
    if not summary:
        return

    doc.add_heading("AI综合分析", level=1)
    doc.add_heading("整体态势评估", level=2)
    overall = summary.get("overall_assessment", "")
    _styled(doc, overall, size=SZ_NORMAL, color=CLR_BODY)

    # 原告最强论点
    p_args = summary.get("plaintiff_strongest_arguments", [])
    if p_args:
        _styled(doc, "原告最强论点", bold=True, size=SZ_SECTION_HDR, color=CLR_BLUE)
        for a in p_args:
            p = doc.add_paragraph()
            _add_run(p, f"[{_h(a['issue_id'])}] ", bold=True, size=SZ_NORMAL, color=CLR_BLUE)
            _add_run(p, a.get("position", ""), size=SZ_NORMAL, color=CLR_BODY)
            reasoning = a.get("reasoning", "")
            if reasoning:
                _styled(doc, f"▶ {reasoning}", size=SZ_NORMAL, color=CLR_GRAY)

    # NOTE: 被告最强抗辩 moved to _render_opponent_strategy_warning() to avoid duplication.

    # 未解决争点
    unresolved = summary.get("unresolved_issues", [])
    if unresolved:
        _styled(doc, "关键未解决争点", bold=True, size=SZ_SECTION_HDR, color=CLR_BLUE)
        for u in unresolved:
            iid = u.get("issue_id", "")
            title = u.get("issue_title", "")
            why = u.get("why_unresolved", "")
            _bullet(doc, f"{_h(iid)} {title}：{why}", size=SZ_NORMAL, color=CLR_BODY)


def _render_missing_evidence(doc, result: dict, party_zh: dict[str, str]):
    """证据缺失报告。"""
    missing = result.get("missing_evidence_report", [])
    if not missing:
        return
    doc.add_heading("证据缺失报告", level=1)
    for m in missing:
        raw_party = m.get("missing_for_party_id", "")
        party_name = party_zh.get(raw_party, raw_party)
        p = doc.add_paragraph()
        _add_run(
            p,
            f"[{_h(m['issue_id'])}] {party_name}: {m['description']}",
            bold=True,
            size=SZ_BODY,
            color=CLR_ORANGE,
        )


def _render_issue_ranking(doc, exec_summary: dict):
    """争点影响排序。"""
    doc.add_heading("争点影响排序", level=1)
    top5 = exec_summary.get("top5_decisive_issues", [])
    _styled(doc, "前五大决定性争点:", bold=True, size=SZ_SECTION_HDR, color=CLR_BLUE)
    visible = _filter_uuids(top5) if isinstance(top5, list) else []
    if visible:
        for iid in visible:
            _bullet(doc, _h(iid), size=SZ_NORMAL, color=CLR_BODY)
    else:
        _styled(doc, "（暂无排序数据）", size=SZ_NORMAL, color=CLR_GRAY)


_PARTY_ZH = {"plaintiff": "原告", "defendant": "被告", "neutral": "中性"}


def _ordered_paths_for_output(decision_tree: dict) -> list[dict]:
    paths = list(decision_tree.get("paths", []))
    path_ranking = decision_tree.get("path_ranking", [])
    if not path_ranking:
        return paths

    rank_index = {
        item.get("path_id"): idx for idx, item in enumerate(path_ranking) if item.get("path_id")
    }
    if not rank_index:
        return paths

    fallback_base = len(rank_index)
    ordered: list[tuple[int, int, dict]] = []
    for source_index, path in enumerate(paths):
        order_key = rank_index.get(path.get("path_id"), fallback_base + source_index)
        ordered.append((order_key, source_index, path))
    ordered.sort(key=lambda item: (item[0], item[1]))
    return [path for _, _, path in ordered]


def _render_decision_tree_probability_free(doc, decision_tree: dict):
    paths = decision_tree.get("paths", [])
    if not paths:
        doc.add_heading("裁判路径树", level=1)
        _styled(
            doc,
            "（本次运行未生成裁判路径，可能需要重新运行庭后分析流程）",
            size=SZ_NORMAL,
            color=CLR_GRAY,
        )
        return

    ordered_paths = _ordered_paths_for_output(decision_tree)
    doc.add_heading(f"裁判路径树（{len(ordered_paths)}条路径）", level=1)

    for rank, path in enumerate(ordered_paths, start=1):
        party = {"plaintiff": "原告", "defendant": "被告", "neutral": "neutral"}.get(
            path.get("party_favored", "neutral"),
            "neutral",
        )
        label_suffix = f"  【有利方：{party}】" if party != "neutral" else ""
        _styled(doc, f"路径 {rank}{label_suffix}", bold=True, size=SZ_SECTION_HDR, color=CLR_BLUE)

        fields = [
            ("触发条件", path.get("trigger_condition", "")),
            ("触发争点", ", ".join(_h(x) for x in path.get("trigger_issue_ids", []))),
            ("关键证据", ", ".join(_h(x) for x in path.get("key_evidence_ids", []))),
            ("可能结果", path.get("possible_outcome", "")),
        ]
        notes = path.get("path_notes", "")
        if notes:
            fields.append(("备注", notes))

        for field_label, val in fields:
            if not val:
                continue
            p = doc.add_paragraph()
            _add_run(p, f"{field_label}：", bold=True, size=SZ_RISK, color=CLR_GRAY)
            _add_run(p, val, size=SZ_RISK, color=CLR_BODY)
        doc.add_paragraph()

    blocking = decision_tree.get("blocking_conditions", [])
    if blocking:
        _styled(doc, "阻断条件", bold=True, size=SZ_SECTION_HDR, color=CLR_RED)
        for bc in blocking:
            _bullet(
                doc,
                f"{_h(bc['condition_id'])}: {bc['description']}",
                size=SZ_NORMAL,
                color=CLR_BODY,
            )


def _render_action_recommendations(doc, exec_summary: dict):
    """行动建议。"""
    doc.add_heading("行动建议", level=1)

    stable = exec_summary.get("current_most_stable_claim", "")
    _styled(doc, "最稳诉请版本：", bold=True, size=SZ_SECTION_HDR, color=CLR_GREEN)
    _styled(doc, stable if stable else "（暂无稳定诉请版本）", size=SZ_NORMAL, color=CLR_BODY)

    actions = exec_summary.get("top3_immediate_actions", [])
    if actions and actions != "未启用":
        _styled(doc, "前三项立即行动：", bold=True, size=SZ_SECTION_HDR, color=CLR_ORANGE)
        if isinstance(actions, list):
            for a in _filter_uuids(actions):
                _bullet(doc, _h(a), size=SZ_NORMAL, color=CLR_BODY)
        else:
            _styled(doc, str(actions), size=SZ_NORMAL, color=CLR_GRAY)

    gaps = exec_summary.get("critical_evidence_gaps", [])
    if gaps and gaps != "未启用":
        _styled(doc, "关键缺证：", bold=True, size=SZ_SECTION_HDR, color=CLR_ORANGE)
        if isinstance(gaps, list):
            for g in _filter_uuids(gaps):
                _bullet(doc, _h(g), size=SZ_NORMAL, color=CLR_BODY)
        else:
            _styled(doc, str(gaps), size=SZ_NORMAL, color=CLR_GRAY)


def _render_executive_summary(doc, exec_summary: dict, amount_report: dict):
    """执行摘要。"""
    doc.add_heading("执行摘要", level=1)

    sections = [
        ("前五大决定性争点", exec_summary.get("top5_decisive_issues", [])),
        ("前三项对方最优攻击", exec_summary.get("top3_adversary_optimal_attacks", [])),
    ]
    for label, val in sections:
        p = doc.add_paragraph()
        _add_run(p, f"{label}：", bold=True, size=SZ_SECTION_HDR, color=CLR_BLUE)
        if isinstance(val, list):
            for item in val:
                _bullet(doc, _h(item), size=SZ_NORMAL, color=CLR_BODY)
        else:
            _styled(doc, str(val), size=SZ_NORMAL, color=CLR_BODY)

    # 金额一致性校验
    check = amount_report.get("consistency_check_result", {})
    verdict_block = check.get("verdict_block_active", False)
    n_conflicts = len(check.get("unresolved_conflicts", []))
    p = doc.add_paragraph()
    _add_run(p, "金额一致性校验：", bold=True, size=SZ_SECTION_HDR, color=CLR_BLUE)
    _add_run(
        p,
        f"阻断裁判={'是' if verdict_block else '否'}，未解决冲突={n_conflicts}条",
        size=SZ_NORMAL,
        color=CLR_RED if verdict_block else CLR_GREEN,
    )


# ---------------------------------------------------------------------------
# Unit 11: 报告增强 — 新增章节渲染函数
# ---------------------------------------------------------------------------

_RISK_COLOR: dict[str, RGBColor] = {
    "favorable": CLR_GREEN,
    "neutral": CLR_ORANGE,
    "unfavorable": CLR_RED,
}


def _render_action_priority_list(doc, exec_summary: dict | None, action_rec: Any | None):
    """行动优先级清单 — 你现在最该做的 3 件事。"""
    items: list[str] = []

    if exec_summary and isinstance(exec_summary.get("top3_immediate_actions"), list):
        items = exec_summary["top3_immediate_actions"][:3]
    elif action_rec is not None:
        # Fallback: derive from action_rec fields
        ev_prios = getattr(action_rec, "evidence_supplement_priorities", None)
        if ev_prios:
            items.append(f"补强证据: {ev_prios[0]}")
        amendments = getattr(action_rec, "recommended_claim_amendments", None)
        if amendments:
            items.append(f"调整诉请: {amendments[0].amendment_description}")
        abandons = getattr(action_rec, "claims_to_abandon", None)
        if abandons:
            items.append(f"考虑放弃: {abandons[0].abandon_reason}")

    if not items:
        return

    doc.add_heading("你现在最该做的 3 件事", level=1)
    visible = _filter_uuids(items)
    for i, item in enumerate(visible[:3], 1):
        p = doc.add_paragraph()
        _add_run(p, f"{i}. ", bold=True, size=SZ_SECTION_HDR, color=CLR_ORANGE)
        _add_run(p, _h(item), size=SZ_BODY, color=CLR_BODY)
    doc.add_paragraph()


def _render_risk_heatmap(doc, ranked_issues: Any | None):
    """风险热力图表格。"""
    rows = build_risk_heatmap(ranked_issues)
    if not rows:
        return

    doc.add_heading(f"风险热力图（{len(rows)}个争点）", level=1)

    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "争点"
    hdr[1].text = "结果影响"
    hdr[2].text = "攻击强度"
    hdr[3].text = "证据强度"
    hdr[4].text = "风险等级"
    hdr[5].text = "建议行动"

    for row in rows:
        cells = table.add_row().cells
        cells[0].text = f"{_h(row.issue_id)}: {row.title[:20]}"
        cells[1].text = row.outcome_impact or "-"
        cells[2].text = row.attack_strength or "-"
        cells[3].text = row.evidence_strength or "-"
        label = RISK_LABEL_ZH.get(row.risk_level, "")
        cells[4].text = label
        # Color the risk cell
        color = _RISK_COLOR.get(row.risk_level.value, CLR_BODY)
        for para in cells[4].paragraphs:
            for run in para.runs:
                run.font.color.rgb = color
                run.font.bold = True
        cells[5].text = _h(row.recommended_action) or "-"

    _set_table_font(table)


def _render_mediation_range(doc, amount_report: dict | None, decision_tree: dict | None):
    """调解区间评估。"""
    med = compute_mediation_range(amount_report, decision_tree)
    if med is None:
        return

    doc.add_heading("调解区间评估", level=1)

    rows_data = [
        ("诉请总额", f"{med.total_claimed:,} 元"),
        ("可核实金额", f"{med.total_verified:,} 元"),
        ("最低可能", f"{med.min_amount:,} 元"),
        ("最高可能", f"{med.max_amount:,} 元"),
        ("建议调解点", f"{med.suggested_amount:,} 元"),
    ]

    table = doc.add_table(rows=len(rows_data), cols=2)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, val) in enumerate(rows_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = _h(val)
    _set_table_font(table)

    # Highlight the suggested amount row
    last_row = table.rows[-1]
    for cell in last_row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = CLR_GREEN

    _styled(doc, f"计算依据: {med.rationale}", size=SZ_RISK, color=CLR_GRAY)


def _render_opponent_strategy_warning(doc, result: dict, attack_chain: dict | None):
    """对方策略预警 — unified section combining defendant defenses and attack chain.

    This is the SINGLE place that renders both ``defendant_strongest_defenses``
    (from the LLM summary) and ``top_attacks`` (from the attack chain).  The
    duplicate renderings in ``_render_llm_summary`` and the standalone
    ``_render_attack_chain`` call have been removed.
    """
    summary = result.get("summary")
    defenses = summary.get("defendant_strongest_defenses", []) if summary else []
    ac = attack_chain or {}
    attacks = ac.get("top_attacks", [])

    if not defenses and not attacks:
        return

    doc.add_heading("对方策略预警", level=1)

    # --- Part 1: 被告核心抗辩及应对建议 ---
    if defenses:
        _styled(doc, "被告核心抗辩及应对建议", bold=True, size=SZ_SECTION_HDR, color=CLR_RED)
        for d in defenses:
            p = doc.add_paragraph()
            _add_run(p, f"[{_h(d['issue_id'])}] ", bold=True, size=SZ_NORMAL, color=CLR_RED)
            _add_run(p, d.get("position", ""), size=SZ_NORMAL, color=CLR_BODY)
            reasoning = d.get("reasoning", "")
            if reasoning:
                _styled(doc, f"对方论据: {reasoning}", size=SZ_RISK, color=CLR_GRAY)
            # Match counter_measure from attack_chain
            if attacks:
                target_id = d.get("issue_id", "")
                for node in attacks:
                    if node.get("target_issue_id") == target_id:
                        cm = node.get("counter_measure", "")
                        if cm:
                            _styled(doc, f"应对建议: {cm}", size=SZ_RISK, color=CLR_GREEN)
                        break
        doc.add_paragraph()

    # --- Part 2: 对方最优攻击路径预警 (replaces standalone _render_attack_chain) ---
    if attacks:
        _styled(doc, "对方最优攻击路径预警", bold=True, size=SZ_SECTION_HDR, color=CLR_RED)

        # Attack chain metadata (previously only in _render_attack_chain)
        raw_owner = ac.get("owner_party_id", "")
        if raw_owner:
            p = doc.add_paragraph()
            _add_run(p, "攻击方：", bold=True, size=SZ_BODY, color=CLR_RED)
            _add_run(p, _h(raw_owner), size=SZ_BODY, color=CLR_BODY)
        order = ac.get("recommended_order", [])
        if order:
            p = doc.add_paragraph()
            _add_run(p, "推荐攻击顺序：", bold=True, size=SZ_BODY, color=CLR_RED)
            _add_run(p, " → ".join(_h(x) for x in order), size=SZ_BODY, color=CLR_BODY)
        doc.add_paragraph()

        for node in attacks:
            nid = node.get("attack_node_id", "")
            target = node.get("target_issue_id", "")
            desc = node.get("attack_description", "")
            p = doc.add_paragraph()
            _add_run(p, f"{_h(nid)} → {_h(target)}: ", bold=True, size=SZ_RISK, color=CLR_RED)
            _add_run(p, desc, size=SZ_RISK, color=CLR_BODY)

            cond = node.get("success_conditions", "")
            if cond:
                _styled(doc, f"成功条件: {cond}", size=SZ_RISK, color=CLR_GRAY)
            evidence = ", ".join(_h(x) for x in node.get("supporting_evidence_ids", []))
            if evidence:
                _styled(doc, f"支撑证据: {evidence}", size=SZ_RISK, color=CLR_GRAY)
            cm = node.get("counter_measure", "")
            if cm:
                _styled(doc, f"应对: {cm}", size=SZ_RISK, color=CLR_GREEN)
            pivot = node.get("adversary_pivot_strategy", "")
            if pivot:
                _styled(doc, f"对方可能转向: {pivot}", size=SZ_RISK, color=CLR_ORANGE)


# ---------------------------------------------------------------------------
# 文书草稿章节 / Document draft section renderers
# ---------------------------------------------------------------------------

_DOC_TYPE_ZH: dict[str, str] = {
    "pleading": "起诉状草稿",
    "defense": "答辩状草稿",
    "cross_exam": "质证意见草稿",
}


def _render_document_drafts(doc, document_drafts: list) -> None:
    """文书草稿章节 — 将 DocumentDraft 列表渲染为 DOCX 结构化章节。
    Document draft sections — renders a list of DocumentDraft objects into DOCX.

    每份文书草稿为独立一级标题，骨架字段按语义分组渲染。
    Each document draft gets its own top-level heading, skeleton fields rendered by group.
    """
    doc.add_heading("文书草稿", level=1)
    _styled(
        doc,
        "以下文书草稿由 AI 生成，仅供律师参考和编辑，不构成正式法律文件。",
        size=SZ_NORMAL,
        color=CLR_ORANGE,
    )
    doc.add_paragraph()

    for draft in document_drafts:
        doc_type = getattr(draft, "doc_type", "")
        case_type = getattr(draft, "case_type", "")
        heading_text = _DOC_TYPE_ZH.get(doc_type, doc_type)
        doc.add_heading(f"{heading_text}（{case_type}）", level=2)

        content = getattr(draft, "content", None)
        if content is None:
            _styled(doc, "（无内容）", size=SZ_NORMAL, color=CLR_GRAY)
            continue

        # 文档头
        header = getattr(content, "header", "")
        if header:
            _styled(doc, header, bold=True, size=SZ_SECTION_HDR, color=CLR_TITLE_DARK)

        # 按文书类型渲染各骨架字段
        if doc_type == "pleading":
            _render_list_section(doc, "事实陈述", getattr(content, "fact_narrative_items", []))
            _render_list_section(doc, "法律依据", getattr(content, "legal_claim_items", []))
            _render_list_section(doc, "诉讼请求", getattr(content, "prayer_for_relief_items", []))
            attack_basis = getattr(content, "attack_chain_basis", "unavailable")
            if attack_basis and attack_basis != "unavailable":
                _styled(doc, f"攻击链策略依据：{attack_basis}", size=SZ_NORMAL, color=CLR_BLUE)

        elif doc_type == "defense":
            _render_list_section(doc, "逐项否认", getattr(content, "denial_items", []))
            _render_list_section(doc, "实质性抗辩", getattr(content, "defense_claim_items", []))
            _render_list_section(doc, "反请求", getattr(content, "counter_prayer_items", []))

        elif doc_type == "cross_exam":
            items = getattr(content, "items", [])
            if items:
                doc.add_heading("逐证据质证意见", level=3)
                for item in items:
                    ev_id = getattr(item, "evidence_id", "")
                    opinion = getattr(item, "opinion_text", "")
                    p = doc.add_paragraph()
                    _add_run(p, f"[{_h(ev_id)}] ", bold=True, size=SZ_NORMAL, color=CLR_BLUE)
                    _add_run(p, opinion, size=SZ_NORMAL, color=CLR_BODY)
            else:
                _styled(doc, "（无证据需要质证）", size=SZ_NORMAL, color=CLR_GRAY)

        # 证据引用
        ev_cited = getattr(draft, "evidence_ids_cited", [])
        if ev_cited:
            p = doc.add_paragraph()
            _add_run(p, "引用证据：", bold=True, size=SZ_NORMAL, color=CLR_GRAY)
            _add_run(p, ", ".join(_h(x) for x in ev_cited), size=SZ_NORMAL, color=CLR_GRAY)

        doc.add_paragraph()


def _render_list_section(doc, label: str, items: list) -> None:
    """渲染一个有序列表章节（标签 + 条目列表）。

    items may be list[str] (legacy) or list[NumberedItem] (seq+text).
    """
    if not items:
        return
    doc.add_heading(label, level=3)
    for item in items:
        text = getattr(item, "text", str(item))
        _bullet(doc, text, size=SZ_NORMAL, color=CLR_BODY)


def _render_similar_cases(doc, similar_cases: list) -> None:
    """类案检索参考章节。

    Args:
        similar_cases: RankedCase 序列化 dict 列表（model_dump 输出）
    """
    _styled(doc, "类案检索参考", bold=True, size=SZ_SUBTITLE, color=CLR_TITLE_DARK)
    _styled(
        doc,
        "以下案例来源于人民法院案例库（rmfyalk.court.gov.cn），经最高人民法院审核认可，"
        "按与本案相关性由高到低排列。",
        size=SZ_NORMAL,
        color=CLR_GRAY,
    )
    doc.add_paragraph()

    for idx, rc in enumerate(similar_cases, 1):
        case_info = rc.get("case", rc) if isinstance(rc, dict) else rc
        relevance = rc.get("relevance", {}) if isinstance(rc, dict) else {}
        analysis = rc.get("analysis", "") if isinstance(rc, dict) else ""

        if isinstance(case_info, dict):
            case_number = case_info.get("case_number", "")
            court = case_info.get("court", "")
            cause = case_info.get("cause_of_action", "")
            kw_list = case_info.get("keywords", [])
            summary = case_info.get("summary", "")
            url = case_info.get("url", "")
        else:
            case_number = getattr(case_info, "case_number", "")
            court = getattr(case_info, "court", "")
            cause = getattr(case_info, "cause_of_action", "")
            kw_list = getattr(case_info, "keywords", [])
            summary = getattr(case_info, "summary", "")
            url = getattr(case_info, "url", "")

        if isinstance(relevance, dict):
            overall = relevance.get("overall", 0)
        else:
            overall = getattr(relevance, "overall", 0)

        # 标题行
        p_title = doc.add_paragraph()
        _add_run(
            p_title,
            f"类案 {idx}：{case_number}",
            bold=True,
            size=SZ_SECTION_HDR,
            color=CLR_BLUE,
        )

        # 基本信息表格
        table = doc.add_table(rows=5, cols=2, style="Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        fields = [
            ("审理法院", court),
            ("案由", cause),
            ("关键词", "、".join(kw_list) if kw_list else "—"),
            ("裁判要旨", summary or "—"),
            ("相关性评分", f"{overall:.0%}"),
        ]
        for row_idx, (label, value) in enumerate(fields):
            table.cell(row_idx, 0).text = label
            table.cell(row_idx, 1).text = str(value)
        _set_table_font(table)

        # 相关性分析
        if analysis:
            p_analysis = doc.add_paragraph()
            _add_run(p_analysis, "与本案相关性分析：", bold=True, size=SZ_NORMAL, color=CLR_BODY)
            _styled(doc, str(analysis), size=SZ_NORMAL, color=CLR_BODY)

        # 案例库链接
        if url:
            _styled(doc, f"案例库链接：{url}", size=SZ_EVIDENCE, color=CLR_GRAY)

        doc.add_paragraph()  # 案例间间隔


# ---------------------------------------------------------------------------
# V3 四层报告 DOCX 生成器 / V3 4-Layer Report DOCX Generator
# ---------------------------------------------------------------------------

_TRAFFIC_LIGHT_EMOJI: dict[str, str] = {
    "green": "🟢",
    "yellow": "🟡",
    "red": "🔴",
}


def generate_docx_v3_report(
    *,
    output_dir: Path,
    report_v3: dict,
    similar_cases: list | None = None,
    filename: str | None = None,
) -> Path:
    """生成 V3 四层架构 Word 报告。

    Args:
        output_dir:    输出目录
        report_v3:     FourLayerReport 序列化 dict（来自 report_v3.json）
        similar_cases: RankedCase 序列化 dict 列表（可选）
        filename:      自定义文件名（默认 "对抗分析报告_v3.docx"）

    Returns:
        生成的 docx 文件路径
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Emu(SZ_NORMAL)

    layer1 = report_v3.get("layer1", {})
    layer2 = report_v3.get("layer2", {})
    layer3 = report_v3.get("layer3", {})
    layer4 = report_v3.get("layer4", {})
    perspective = report_v3.get("perspective", "neutral")

    _render_v3_title(doc, report_v3)
    _render_disclaimer(doc)
    _render_v3_layer1(doc, layer1, perspective)
    _render_v3_layer2(doc, layer2)
    _render_v3_layer3(doc, layer3, perspective)
    _render_v3_layer4(doc, layer4)
    if similar_cases:
        _render_similar_cases(doc, similar_cases)

    if filename is None:
        filename = "对抗分析报告_v3.docx"
    dest = output_dir / filename
    doc.save(str(dest))
    return dest


def _render_v3_title(doc, report_v3: dict) -> None:
    """V3 封面标题块。"""
    case_id = report_v3.get("case_id", "")
    run_id = report_v3.get("run_id", "")
    perspective = report_v3.get("perspective", "neutral")
    perspective_label = {
        "plaintiff": "原告视角",
        "defendant": "被告视角",
        "neutral": "中立双视角",
    }.get(perspective, "中立双视角")

    p0 = doc.add_paragraph()
    _add_run(p0, "对抗分析报告（V3 四层架构）", bold=True, size=SZ_TITLE, color=CLR_TITLE_DARK)

    p1 = doc.add_paragraph()
    _add_run(p1, f"Case ID: {case_id}  |  Run ID: {run_id}", size=SZ_NORMAL, color=CLR_GRAY)

    p2 = doc.add_paragraph()
    _add_run(
        p2, f"视角: {perspective_label}  |  报告版本: V3 四层架构", size=SZ_NORMAL, color=CLR_GRAY
    )
    doc.add_paragraph()


def _render_v3_layer1(doc, layer1: dict, perspective: str) -> None:
    """一、封面摘要 「事实」 — supports V3.1 fields with V3.0 fallback."""
    doc.add_heading("一、封面摘要 「事实」", level=1)

    cover_summary = layer1.get("cover_summary", {})

    # A. 中立结论摘要 (unchanged across versions)
    doc.add_heading("A. 中立结论摘要 「推断」", level=2)
    neutral = cover_summary.get("neutral_conclusion", "")
    if neutral:
        p = doc.add_paragraph()
        _add_run(p, neutral, bold=True, size=SZ_BODY, color=CLR_BODY)
    doc.add_paragraph()

    # B. 胜负手 (V3.1) — fallback to plaintiff/defendant summaries (V3.0)
    winning_move = cover_summary.get("winning_move") or ""
    if winning_move:
        doc.add_heading("B. 胜负手 「建议」", level=2)
        p = doc.add_paragraph()
        _add_run(p, winning_move, bold=True, size=SZ_BODY, color=CLR_RED)
        doc.add_paragraph()
    else:
        # V3.0 fallback: plaintiff + defendant summaries
        plaintiff_summary = cover_summary.get("plaintiff_summary") or {}
        if plaintiff_summary:
            doc.add_heading("B-1. 原告视角 「建议」", level=2)
            for i, s in enumerate(plaintiff_summary.get("top3_strengths", []), 1):
                p = doc.add_paragraph()
                _add_run(p, f"优势{i}: ", bold=True, size=SZ_BODY, color=CLR_GREEN)
                _add_run(p, s, size=SZ_BODY, color=CLR_BODY)
            for i, d in enumerate(plaintiff_summary.get("top2_dangers", []), 1):
                p = doc.add_paragraph()
                _add_run(p, f"危险{i}: ", bold=True, size=SZ_BODY, color=CLR_RED)
                _add_run(p, d, size=SZ_BODY, color=CLR_BODY)
            doc.add_paragraph()

        defendant_summary = cover_summary.get("defendant_summary") or {}
        if defendant_summary:
            doc.add_heading("B-2. 被告视角 「建议」", level=2)
            for i, d in enumerate(defendant_summary.get("top3_defenses", []), 1):
                p = doc.add_paragraph()
                _add_run(p, f"防线{i}: ", bold=True, size=SZ_BODY, color=CLR_BLUE)
                _add_run(p, d, size=SZ_BODY, color=CLR_BODY)
            visible_supp = _filter_uuids(defendant_summary.get("plaintiff_likely_supplement", []))
            if visible_supp:
                _styled(doc, "原告可能补强方向：", bold=True, size=SZ_RISK, color=CLR_ORANGE)
                for s in visible_supp:
                    _bullet(doc, _h(s), size=SZ_RISK, color=CLR_BODY)
            visible_order = _filter_uuids(defendant_summary.get("optimal_attack_order", []))
            if visible_order:
                p = doc.add_paragraph()
                _add_run(p, "最优攻击顺序：", bold=True, size=SZ_RISK, color=CLR_RED)
                _add_run(p, " → ".join(_h(x) for x in visible_order), size=SZ_RISK, color=CLR_BODY)
            doc.add_paragraph()

    # C. 阻断条件 (V3.1) — fallback to scenario_tree_summary (V3.0)
    blocking_conditions = cover_summary.get("blocking_conditions") or []
    if blocking_conditions:
        doc.add_heading("C. 阻断条件 「推断」", level=2)
        for i, cond in enumerate(blocking_conditions, 1):
            p = doc.add_paragraph()
            _add_run(p, f"{i}. ", bold=True, size=SZ_BODY, color=CLR_ORANGE)
            _add_run(p, cond, size=SZ_BODY, color=CLR_BODY)
        doc.add_paragraph()
    else:
        scenario_summary = layer1.get("scenario_tree_summary", "")
        if scenario_summary:
            doc.add_heading("C. 条件场景摘要 「推断」", level=2)
            conditions = [
                c.strip() for c in scenario_summary.replace("；", ";").split(";") if c.strip()
            ]
            for cond in conditions:
                _bullet(doc, cond, size=SZ_BODY, color=CLR_BODY)
            doc.add_paragraph()

    # D. 案件时间线 (V3.1)
    timeline = layer1.get("timeline") or []
    if timeline:
        doc.add_heading("D. 案件时间线 「事实」", level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        hdr[0].text = "日期"
        hdr[1].text = "事件"
        hdr[2].text = "来源"
        hdr[3].text = "争议"
        for evt in timeline:
            if isinstance(evt, dict):
                date = evt.get("date", "")
                event = evt.get("event", "")
                source = evt.get("source", "")
                disputed = evt.get("disputed", False)
            else:
                date = getattr(evt, "date", "")
                event = getattr(evt, "event", "")
                source = getattr(evt, "source", "")
                disputed = getattr(evt, "disputed", False)
            row = table.add_row().cells
            row[0].text = date
            row[1].text = event
            row[2].text = _h(source)
            row[3].text = "⚠ 有争议" if disputed else ""
        _set_table_font(table)
        doc.add_paragraph()

    # E. 证据优先级 (V3.1) — fallback to evidence_traffic_lights (V3.0)
    evidence_priorities = layer1.get("evidence_priorities") or []
    if evidence_priorities:
        doc.add_heading("E. 证据优先级 「事实」", level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        hdr[0].text = "证据"
        hdr[1].text = "优先级"
        hdr[2].text = "理由"
        for ep in evidence_priorities:
            if isinstance(ep, dict):
                ev_id = ep.get("evidence_id", "")
                title = ep.get("title", "")
                priority = ep.get("priority", "")
                reason = ep.get("reason", "")
            else:
                ev_id = getattr(ep, "evidence_id", "")
                title = getattr(ep, "title", "")
                priority = getattr(ep, "priority", "")
                reason = getattr(ep, "reason", "")
            label = f"{_h(ev_id)}: {title}" if title else _h(ev_id)
            row = table.add_row().cells
            row[0].text = label
            row[1].text = priority
            row[2].text = reason
        _set_table_font(table)
        doc.add_paragraph()
    else:
        # V3.0 fallback: evidence traffic lights
        traffic_lights = layer1.get("evidence_traffic_lights", [])
        if traffic_lights:
            doc.add_heading("D. 证据风险红绿灯 「事实」", level=2)
            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = table.rows[0].cells
            hdr[0].text = "证据 ID"
            hdr[1].text = "标题"
            hdr[2].text = "风险"
            hdr[3].text = "理由"
            for tl in traffic_lights:
                if isinstance(tl, dict):
                    ev_id = tl.get("evidence_id", "")
                    title = tl.get("title", "")
                    risk = str(tl.get("risk_level", "yellow"))
                    reason = tl.get("reason", "")
                else:
                    ev_id = getattr(tl, "evidence_id", "")
                    title = getattr(tl, "title", "")
                    risk = str(getattr(tl, "risk_level", "yellow"))
                    reason = getattr(tl, "reason", "")
                row = table.add_row().cells
                row[0].text = _h(ev_id)
                row[1].text = title
                row[2].text = _TRAFFIC_LIGHT_EMOJI.get(risk, "🟡")
                row[3].text = reason
            _set_table_font(table)
            doc.add_paragraph()


def _render_v3_layer2(doc, layer2: dict) -> None:
    """二、中立对抗内核 「事实」 — supports V3.1 fields with V3.0 fallback."""
    doc.add_heading("二、中立对抗内核 「事实」", level=1)

    # 2.1 事实底座 (unchanged across versions)
    doc.add_heading("2.1 事实底座 「事实」", level=2)
    fact_base = layer2.get("fact_base", [])
    if fact_base:
        for entry in fact_base:
            if isinstance(entry, dict):
                desc = entry.get("description", "")
                sources = entry.get("source_evidence_ids", [])
            else:
                desc = getattr(entry, "description", "")
                sources = getattr(entry, "source_evidence_ids", [])
            p = doc.add_paragraph()
            _add_run(p, "• ", bold=True, size=SZ_BODY, color=CLR_BLUE)
            _add_run(p, desc, size=SZ_BODY, color=CLR_BODY)
            if sources:
                _styled(
                    doc,
                    f"  来源证据: {', '.join(_h(x) for x in sources)}",
                    size=SZ_EVIDENCE,
                    color=CLR_GRAY,
                )
    else:
        _styled(doc, "暂无双方均认可的无争议事实。", size=SZ_NORMAL, color=CLR_GRAY)
    doc.add_paragraph()

    # 2.2 争点地图 — V3.1 depth-aware rendering
    doc.add_heading("2.2 争点地图 「推断」", level=2)
    issue_map = layer2.get("issue_map", [])
    for card in issue_map:
        if isinstance(card, dict):
            issue_id = card.get("issue_id", "")
            issue_title = card.get("issue_title", "")
            depth = card.get("depth", 0)
            plaintiff_thesis = card.get("plaintiff_thesis", "")
            defendant_thesis = card.get("defendant_thesis", "")
            decisive_evidence = card.get("decisive_evidence", [])
            current_gaps = card.get("current_gaps", [])
            outcome_sensitivity = card.get("outcome_sensitivity", "")
        else:
            issue_id = getattr(card, "issue_id", "")
            issue_title = getattr(card, "issue_title", "")
            depth = getattr(card, "depth", 0)
            plaintiff_thesis = getattr(card, "plaintiff_thesis", "")
            defendant_thesis = getattr(card, "defendant_thesis", "")
            decisive_evidence = getattr(card, "decisive_evidence", [])
            current_gaps = getattr(card, "current_gaps", [])
            outcome_sensitivity = getattr(card, "outcome_sensitivity", "")

        # V3.1: depth=0 as level-3 heading, depth>0 as indented paragraph
        if depth == 0:
            doc.add_heading(f"{_h(issue_id)}: {issue_title}", level=3)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Emu(depth * 457_200)  # 0.5 inch per level
            _add_run(
                p,
                f"└─ {_h(issue_id)}: {issue_title}",
                bold=True,
                size=SZ_SECTION_HDR,
                color=CLR_BLUE,
            )
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        hdr[0].text = "字段"
        hdr[1].text = "内容"
        fields = [
            ("原告主张", plaintiff_thesis[:500] + ("..." if len(plaintiff_thesis) > 500 else "")),
            ("被告主张", defendant_thesis[:500] + ("..." if len(defendant_thesis) > 500 else "")),
            ("决定性证据", ", ".join(_h(x) for x in decisive_evidence)),
            ("当前缺口", " | ".join(current_gaps) if current_gaps else "暂无"),
            ("结果敏感度", outcome_sensitivity),
        ]
        for label, value in fields:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = value or "—"
        _set_table_font(table)
        doc.add_paragraph()

    # 2.3 证据卡片 (V3.1) — fallback to evidence_battle_matrix (V3.0)
    evidence_cards = layer2.get("evidence_cards") or []
    if evidence_cards:
        # V3.1: unified electronic evidence strategy
        unified_strategy = layer2.get("unified_electronic_strategy") or ""
        if unified_strategy:
            doc.add_heading("2.3 统一电子证据补强策略 「建议」", level=2)
            p = doc.add_paragraph()
            _add_run(p, unified_strategy, size=SZ_BODY, color=CLR_BODY)
            doc.add_paragraph()

        doc.add_heading("2.4 证据卡片 「推断」", level=2)
        for card in evidence_cards:
            if isinstance(card, dict):
                ev_id = card.get("evidence_id", "")
                q1 = card.get("q1_what", "")
                q2 = card.get("q2_target", "")
                q3 = card.get("q3_key_risk", "")
                q4 = card.get("q4_best_attack", "")
                q5 = card.get("q5_reinforce", "")
                q6 = card.get("q6_failure_impact", "")
                priority = card.get("priority", "")
            else:
                ev_id = getattr(card, "evidence_id", "")
                q1 = getattr(card, "q1_what", "")
                q2 = getattr(card, "q2_target", "")
                q3 = getattr(card, "q3_key_risk", "")
                q4 = getattr(card, "q4_best_attack", "")
                q5 = getattr(card, "q5_reinforce", "")
                q6 = getattr(card, "q6_failure_impact", "")
                priority = getattr(card, "priority", "")

            is_core = bool(q5)  # Core evidence has q5_reinforce populated

            if is_core:
                # Full 6-field table card for core evidence
                p = doc.add_paragraph()
                _add_run(p, f"{_h(ev_id)}", bold=True, size=SZ_SECTION_HDR, color=CLR_BLUE)
                if priority:
                    _add_run(p, f"  [{priority}]", bold=True, size=SZ_RISK, color=CLR_RED)
                table = doc.add_table(rows=6, cols=2)
                table.style = "Light Grid Accent 1"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                questions = [
                    ("1. 这是什么证据", q1),
                    ("2. 证明目标", q2),
                    ("3. 关键风险", q3),
                    ("4. 最佳攻击方式", q4),
                    ("5. 如何补强", q5),
                    ("6. 失败影响", q6),
                ]
                for i, (label, val) in enumerate(questions):
                    table.rows[i].cells[0].text = label
                    table.rows[i].cells[1].text = _h(val) or "—"
                _set_table_font(table)
                doc.add_paragraph()
            else:
                # Compact row for supporting / background evidence
                p = doc.add_paragraph()
                _add_run(p, f"{_h(ev_id)}", bold=True, size=SZ_BODY, color=CLR_GRAY)
                if priority:
                    _add_run(p, f"  [{priority}]", size=SZ_RISK, color=CLR_ORANGE)
                table = doc.add_table(rows=4, cols=2)
                table.style = "Light Grid Accent 1"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                questions = [
                    ("1. 这是什么证据", q1),
                    ("2. 证明目标", q2),
                    ("3. 关键风险", q3),
                    ("4. 最佳攻击方式", q4),
                ]
                for i, (label, val) in enumerate(questions):
                    table.rows[i].cells[0].text = label
                    table.rows[i].cells[1].text = _h(val) or "—"
                _set_table_font(table)
                doc.add_paragraph()
    else:
        # V3.0 fallback: evidence battle matrix
        evidence_battle_matrix = layer2.get("evidence_battle_matrix", [])
        if evidence_battle_matrix:
            doc.add_heading("2.3 证据作战矩阵 「推断」", level=2)
            for card in evidence_battle_matrix:
                if isinstance(card, dict):
                    ev_id = card.get("evidence_id", "")
                    risk = str(card.get("risk_level", "yellow"))
                    q1 = card.get("q1_what", "")
                    q2 = card.get("q2_proves", "")
                    q3 = card.get("q3_direction", "")
                    q4 = card.get("q4_risks", "")
                    q5 = card.get("q5_opponent_attack", "")
                    q6 = card.get("q6_reinforce", "")
                    q7 = card.get("q7_failure_impact", "")
                else:
                    ev_id = getattr(card, "evidence_id", "")
                    risk = str(getattr(card, "risk_level", "yellow"))
                    q1 = getattr(card, "q1_what", "")
                    q2 = getattr(card, "q2_proves", "")
                    q3 = getattr(card, "q3_direction", "")
                    q4 = getattr(card, "q4_risks", "")
                    q5 = getattr(card, "q5_opponent_attack", "")
                    q6 = getattr(card, "q6_reinforce", "")
                    q7 = getattr(card, "q7_failure_impact", "")

                emoji = _TRAFFIC_LIGHT_EMOJI.get(risk, "🟡")
                p = doc.add_paragraph()
                _add_run(p, f"{_h(ev_id)} {emoji}", bold=True, size=SZ_SECTION_HDR, color=CLR_BLUE)
                table = doc.add_table(rows=7, cols=2)
                table.style = "Light Grid Accent 1"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                questions = [
                    ("1. 这是什么证据", q1),
                    ("2. 证明什么命题", q2),
                    ("3. 证明方向", q3),
                    ("4. 四性风险", q4),
                    ("5. 对方如何攻击", q5),
                    ("6. 如何加固", q6),
                    ("7. 失败影响", q7),
                ]
                for i, (label, val) in enumerate(questions):
                    table.rows[i].cells[0].text = label
                    table.rows[i].cells[1].text = _h(val) or "—"
                _set_table_font(table)
                doc.add_paragraph()


def _render_v3_layer3(doc, layer3: dict, perspective: str) -> None:
    """三、角色化输出 「建议」 — supports V3.1 action sections with V3.0 fallback."""
    doc.add_heading("三、角色化输出 「建议」", level=1)

    outputs = layer3.get("outputs", [])
    for output in outputs:
        if isinstance(output, dict):
            pov = output.get("perspective", "neutral")
            # V3.1 action-oriented fields
            evidence_supplement_checklist = output.get("evidence_supplement_checklist") or []
            cross_examination_points = output.get("cross_examination_points") or []
            trial_questions = output.get("trial_questions") or []
            contingency_plans = output.get("contingency_plans") or []
            over_assertion_boundaries = output.get("over_assertion_boundaries") or []
            unified_electronic_evidence_strategy = (
                output.get("unified_electronic_evidence_strategy") or ""
            )
            # V3.0 legacy fields
            top_claims = output.get("top_claims") or []
            defendant_attack_chains = output.get("defendant_attack_chains") or []
            evidence_to_supplement = output.get("evidence_to_supplement") or []
            trial_sequence = output.get("trial_sequence") or []
            claims_to_abandon = output.get("claims_to_abandon") or []
            top_defenses = output.get("top_defenses") or []
            plaintiff_supplement_prediction = output.get("plaintiff_supplement_prediction") or []
            evidence_to_challenge_first = output.get("evidence_to_challenge_first") or []
            motions_to_file = output.get("motions_to_file") or []
            over_assertion_warnings = output.get("over_assertion_warnings") or []
        else:
            pov = getattr(output, "perspective", "neutral")
            evidence_supplement_checklist = (
                getattr(output, "evidence_supplement_checklist", []) or []
            )
            cross_examination_points = getattr(output, "cross_examination_points", []) or []
            trial_questions = getattr(output, "trial_questions", []) or []
            contingency_plans = getattr(output, "contingency_plans", []) or []
            over_assertion_boundaries = getattr(output, "over_assertion_boundaries", []) or []
            unified_electronic_evidence_strategy = (
                getattr(output, "unified_electronic_evidence_strategy", "") or ""
            )
            top_claims = getattr(output, "top_claims", []) or []
            defendant_attack_chains = getattr(output, "defendant_attack_chains", []) or []
            evidence_to_supplement = getattr(output, "evidence_to_supplement", []) or []
            trial_sequence = getattr(output, "trial_sequence", []) or []
            claims_to_abandon = getattr(output, "claims_to_abandon", []) or []
            top_defenses = getattr(output, "top_defenses", []) or []
            plaintiff_supplement_prediction = (
                getattr(output, "plaintiff_supplement_prediction", []) or []
            )
            evidence_to_challenge_first = getattr(output, "evidence_to_challenge_first", []) or []
            motions_to_file = getattr(output, "motions_to_file", []) or []
            over_assertion_warnings = getattr(output, "over_assertion_warnings", []) or []

        # Detect V3.1: any of the 5 new action fields populated
        has_v31_fields = any(
            [
                evidence_supplement_checklist,
                cross_examination_points,
                trial_questions,
                contingency_plans,
                over_assertion_boundaries,
            ]
        )

        pov_label = {"plaintiff": "原告", "defendant": "被告"}.get(pov, pov)
        doc.add_heading(f"{pov_label}策略 「建议」", level=2)

        if has_v31_fields:
            # V3.1 action-oriented rendering
            if evidence_supplement_checklist:
                doc.add_heading("补证清单 「建议」", level=3)
                for i, item in enumerate(evidence_supplement_checklist, 1):
                    p = doc.add_paragraph()
                    _add_run(p, f"{i}. ", bold=True, size=SZ_BODY, color=CLR_BLUE)
                    _add_run(p, item, size=SZ_BODY, color=CLR_BODY)

            if cross_examination_points:
                doc.add_heading("质证要点 「建议」", level=3)
                for i, item in enumerate(cross_examination_points, 1):
                    p = doc.add_paragraph()
                    _add_run(p, f"{i}. ", bold=True, size=SZ_BODY, color=CLR_RED)
                    _add_run(p, item, size=SZ_BODY, color=CLR_BODY)

            if trial_questions:
                doc.add_heading("庭审发问 「建议」", level=3)
                for i, item in enumerate(trial_questions, 1):
                    p = doc.add_paragraph()
                    _add_run(p, f"{i}. ", bold=True, size=SZ_BODY, color=CLR_BLUE)
                    _add_run(p, item, size=SZ_BODY, color=CLR_BODY)

            if contingency_plans:
                doc.add_heading("应对预案 「建议」", level=3)
                for i, item in enumerate(contingency_plans, 1):
                    p = doc.add_paragraph()
                    _add_run(p, f"{i}. ", bold=True, size=SZ_BODY, color=CLR_ORANGE)
                    _add_run(p, item, size=SZ_BODY, color=CLR_BODY)

            if over_assertion_boundaries:
                doc.add_heading("过度主张边界 「观点」", level=3)
                for item in over_assertion_boundaries:
                    _bullet(doc, item, size=SZ_BODY, color=CLR_RED)

            if unified_electronic_evidence_strategy:
                doc.add_heading("统一电子证据策略 「建议」", level=3)
                p = doc.add_paragraph()
                _add_run(p, unified_electronic_evidence_strategy, size=SZ_BODY, color=CLR_BODY)

            doc.add_paragraph()
        else:
            # V3.0 fallback: role-specific old-style rendering
            if pov == "plaintiff":
                if top_claims:
                    doc.add_heading("三大诉请", level=3)
                    for i, claim in enumerate(top_claims, 1):
                        p = doc.add_paragraph()
                        _add_run(p, f"{i}. ", bold=True, size=SZ_BODY, color=CLR_BLUE)
                        _add_run(p, claim, size=SZ_BODY, color=CLR_BODY)
                if defendant_attack_chains:
                    doc.add_heading("被告攻击链预警 「推断」", level=3)
                    for warning in defendant_attack_chains:
                        _bullet(doc, warning, size=SZ_BODY, color=CLR_ORANGE)
                if evidence_to_supplement:
                    doc.add_heading("需补强证据清单 「建议」", level=3)
                    for i, ev in enumerate(_filter_uuids(evidence_to_supplement), 1):
                        p = doc.add_paragraph()
                        _add_run(p, f"{i}. ", bold=True, size=SZ_NORMAL, color=CLR_GRAY)
                        _add_run(p, ev, size=SZ_NORMAL, color=CLR_BODY)
                if trial_sequence:
                    doc.add_heading("庭审举证顺序建议 「建议」", level=3)
                    for i, step in enumerate(trial_sequence, 1):
                        p = doc.add_paragraph()
                        _add_run(p, f"{i}. ", bold=True, size=SZ_NORMAL, color=CLR_GRAY)
                        _add_run(p, step, size=SZ_NORMAL, color=CLR_BODY)
                if claims_to_abandon:
                    doc.add_heading("应放弃诉请 「建议」", level=3)
                    for item in claims_to_abandon:
                        _bullet(doc, item, size=SZ_NORMAL, color=CLR_RED)
                doc.add_paragraph()

            elif pov == "defendant":
                if top_defenses:
                    doc.add_heading("三大防线", level=3)
                    for i, defense in enumerate(top_defenses, 1):
                        p = doc.add_paragraph()
                        _add_run(p, f"{i}. ", bold=True, size=SZ_BODY, color=CLR_BLUE)
                        _add_run(p, defense, size=SZ_BODY, color=CLR_BODY)
                if plaintiff_supplement_prediction:
                    doc.add_heading("原告可能补强方向 「推断」", level=3)
                    for item in _filter_uuids(plaintiff_supplement_prediction):
                        _bullet(doc, item, size=SZ_NORMAL, color=CLR_ORANGE)
                if evidence_to_challenge_first:
                    doc.add_heading("优先质证目标 「建议」", level=3)
                    for item in evidence_to_challenge_first:
                        _bullet(doc, item, size=SZ_NORMAL, color=CLR_RED)
                if motions_to_file:
                    doc.add_heading("应提交动议 「建议」", level=3)
                    for item in motions_to_file:
                        _bullet(doc, item, size=SZ_NORMAL, color=CLR_BODY)
                if over_assertion_warnings:
                    doc.add_heading("过度主张警告 「观点」", level=3)
                    for item in over_assertion_warnings:
                        _bullet(doc, item, size=SZ_NORMAL, color=CLR_RED)
                doc.add_paragraph()


def _render_v3_layer4(doc, layer4: dict) -> None:
    """四、附录 「事实」"""
    doc.add_heading("四、附录 「事实」", level=1)

    transcripts = layer4.get("adversarial_transcripts_md", "")
    if transcripts:
        doc.add_heading("4.1 三轮对抗辩论记录", level=2)
        _render_md_as_text(doc, _h(transcripts))
        doc.add_paragraph()

    evidence_index = layer4.get("evidence_index_md", "")
    if evidence_index:
        doc.add_heading("4.2 证据索引表", level=2)
        _render_md_as_text(doc, _h(evidence_index))
        doc.add_paragraph()

    timeline = layer4.get("timeline_md", "")
    # V3.1: skip timeline_md if it only contains placeholder text (timeline
    # is now rendered in Layer 1 via structured timeline events)
    _timeline_placeholder = "暂无时间线数据"
    if timeline and _timeline_placeholder not in timeline:
        doc.add_heading("4.3 案件时间线", level=2)
        _render_md_as_text(doc, _h(timeline))
        doc.add_paragraph()

    glossary = layer4.get("glossary_md", "")
    if glossary:
        doc.add_heading("4.4 术语表", level=2)
        _render_md_as_text(doc, _h(glossary))
        doc.add_paragraph()

    amount_calc = layer4.get("amount_calculation_md", "")
    if amount_calc:
        doc.add_heading("4.5 金额计算明细", level=2)
        _render_md_as_text(doc, _h(amount_calc))
        doc.add_paragraph()


def _render_md_as_text(doc, md_text: str) -> None:
    """将 Markdown 文本简单渲染为 DOCX 段落（剥离标记符，保留结构）。"""
    import re as _re

    for line in md_text.splitlines():
        stripped = line.strip()
        if not stripped or stripped == "---":
            doc.add_paragraph()
            continue
        # Headings
        m = _re.match(r"^(#{1,4})\s+(.*)", stripped)
        if m:
            level = min(len(m.group(1)), 4)
            doc.add_heading(m.group(2).strip(), level=level)
            continue
        # Bold lines like **agent_name** — description
        clean = _re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
        clean = _re.sub(r"\*(.+?)\*", r"\1", clean)
        _styled(doc, clean, size=SZ_EVIDENCE, color=CLR_BODY)
