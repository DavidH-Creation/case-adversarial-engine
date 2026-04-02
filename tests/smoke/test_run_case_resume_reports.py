from __future__ import annotations

import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document


FIXTURE_ROOT = Path(__file__).parent / "fixtures" / "resume_post_debate"
RUN_FIXTURE = FIXTURE_ROOT / "run"
CASE_FILE = FIXTURE_ROOT / "case.yaml"
REPO_ROOT = Path(__file__).resolve().parents[2]


def _read_docx_text(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_run_case_resume_regenerates_probability_free_reports(tmp_path: Path) -> None:
    run_dir = tmp_path / "run"
    shutil.copytree(RUN_FIXTURE, run_dir)

    proc = subprocess.run(
        [
            sys.executable,
            "scripts/run_case.py",
            str(CASE_FILE),
            "--output-dir",
            str(run_dir),
            "--resume",
            "--perspective",
            "plaintiff",
        ],
        cwd=REPO_ROOT,
        capture_output=True,
        text=True,
        encoding="utf-8",
    )

    assert proc.returncode == 0, proc.stdout + "\n" + proc.stderr

    report_md = run_dir / "report.md"
    report_docx = run_dir / "report.docx"
    assert report_md.exists()
    assert report_docx.exists()

    md_content = report_md.read_text(encoding="utf-8")
    docx_content = _read_docx_text(report_docx)

    for content in (md_content, docx_content):
        assert "调解区间评估" not in content
        assert "路径概率比较" not in content
        assert "可能性：" not in content
        assert "概率依据" not in content
        assert "prob=" not in content
        assert "置信区间" not in content

    assert "Court orders repayment of the " in md_content
    assert "Court orders repayment of the " in docx_content
    assert "The IOU is authentic and the transfer is proven" in md_content
    assert "The IOU is authentic and the transfer is proven" in docx_content
