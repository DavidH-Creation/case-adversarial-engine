#!/usr/bin/env python3
"""Generate v3 Word document matching v2 format exactly."""

import json
import sys
import os
from pathlib import Path
from decimal import Decimal

sys.path.insert(0, str(Path(__file__).parent.parent))
os.environ.setdefault("PYTHONIOENCODING", "utf-8")

from docx import Document
from docx.shared import Pt, Emu, RGBColor, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# Color palette (matching v2)
# ---------------------------------------------------------------------------
CLR_TITLE_DARK = RGBColor(0x1B, 0x3A, 0x5C)  # dark blue - title
CLR_BLUE = RGBColor(0x2E, 0x75, 0xB6)  # blue - plaintiff / section headers
CLR_RED = RGBColor(0xC0, 0x39, 0x2B)  # red - defendant / weaknesses
CLR_GREEN = RGBColor(0x27, 0xAE, 0x60)  # green - evidence mgr / strengths
CLR_ORANGE = RGBColor(0xE6, 0x7E, 0x22)  # orange - risk flags / warnings
CLR_BODY = RGBColor(0x1A, 0x1A, 0x1A)  # dark - body text
CLR_GRAY = RGBColor(0x4A, 0x4A, 0x4A)  # gray - secondary text

# Font sizes (EMU, matching v2)
SZ_TITLE = 279400  # ~22pt
SZ_SUBTITLE = 203200  # ~16pt
SZ_AGENT_TITLE = 152400  # ~12pt
SZ_SECTION_HDR = 139700  # ~11pt
SZ_BODY = 133350  # ~10.5pt
SZ_RISK = 120650  # ~9.5pt
SZ_EVIDENCE = 114300  # ~9pt
SZ_NORMAL = 127000  # ~10pt
FONT_NAME = "Arial"

# ---------------------------------------------------------------------------
# Translation maps (Chinese localization)
# ---------------------------------------------------------------------------
_ISSUE_TYPE_ZH = {
    "factual": "事实争点",
    "mixed": "混合争点",
    "legal": "法律争点",
    "procedural": "程序争点",
}
_PARTY_ZH = {
    "party-plaintiff-wang": "原告老王方",
    "party-defendant-chen": "被告小陈方",
    "party-defendant-zhuang": "被告老庄方",
}


def _add_run(para, text, *, bold=False, size=None, color=None, italic=False):
    """Add a formatted run to a paragraph."""
    run = para.add_run(text)
    run.font.name = FONT_NAME
    # Set east-asian font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if bold:
        run.font.bold = True
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color
    if italic:
        run.font.italic = True
    return run


def _add_styled_para(doc, text, *, bold=False, size=SZ_NORMAL, color=CLR_BODY):
    """Add a single-run paragraph with consistent styling."""
    p = doc.add_paragraph()
    _add_run(p, text, bold=bold, size=size, color=color)
    return p


def _add_bullet(doc, text, *, size=SZ_NORMAL, color=CLR_BODY):
    """Add a bullet point paragraph."""
    p = doc.add_paragraph()
    _add_run(p, "\u2022 " + text, size=size, color=color)
    return p


def _agent_color(role):
    if "plaintiff" in role:
        return CLR_BLUE
    elif "defendant" in role:
        return CLR_RED
    else:
        return CLR_GREEN


def _set_table_font(table):
    """Set font for all cells in a table."""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = FONT_NAME
                    run.font.size = SZ_NORMAL


# ---------------------------------------------------------------------------
# Load data
# ---------------------------------------------------------------------------
out = Path(__file__).parent.parent / "outputs" / "v3"

result = json.loads((out / "result.json").read_text(encoding="utf-8"))
decision_tree = json.loads((out / "decision_tree.json").read_text(encoding="utf-8"))
attack_chain = json.loads((out / "attack_chain.json").read_text(encoding="utf-8"))
exec_summary = json.loads((out / "exec_summary.json").read_text(encoding="utf-8"))
amount_report = json.loads((out / "amount_report.json").read_text(encoding="utf-8"))

# Parse issue descriptions & types from report.md
import re

_issue_info: dict[str, tuple[str, str]] = {}  # issue_id -> (description, type)
report_md = (out / "report.md").read_text(encoding="utf-8")
for m in re.finditer(r"- \*\*\[([^\]]+)\]\*\*\s+(.+?)\s+`(\w+)`", report_md):
    _issue_info[m.group(1)] = (m.group(2), m.group(3))

doc = Document()

# Set default font
style = doc.styles["Normal"]
style.font.name = FONT_NAME
style.font.size = Emu(SZ_NORMAL)

# ---------------------------------------------------------------------------
# Title block (matching v2: Case XX vN / subtitle / meta line)
# ---------------------------------------------------------------------------
p0 = doc.add_paragraph()
_add_run(p0, "案件09 第三版分析报告", bold=True, size=SZ_TITLE, color=CLR_TITLE_DARK)

p1 = doc.add_paragraph()
_add_run(
    p1,
    "\u501f\u8d37\u5408\u610f\u4e3b\u4f53\u4e89\u8bae\uff08\u8001\u738b\u8bc9\u5c0f\u9648\u3001\u8001\u5e84\uff09",
    size=SZ_SUBTITLE,
    color=CLR_BLUE,
)

p2 = doc.add_paragraph()
_add_run(p2, "对抗分析报告  |  v3 对抗引擎  |  Claude Opus 4.6", size=SZ_NORMAL, color=CLR_GRAY)

doc.add_paragraph()

# ---------------------------------------------------------------------------
# 案件摘要 (Heading 1 + Table)
# ---------------------------------------------------------------------------
doc.add_heading("\u6848\u4ef6\u6458\u8981", level=1)

summary_rows = [
    ("\u6848\u4ef6ID", result["case_id"]),
    ("\u8fd0\u884cID", result["run_id"]),
    ("\u539f\u544a", "\u8001\u738b\uff08\u7537\uff0c\u6c49\u65cf\uff09"),
    (
        "\u88ab\u544a",
        "\u5c0f\u9648\uff08\u5973\uff0c\u6c49\u65cf\uff09\u3001\u8001\u5e84\uff08\u7537\uff0c\u6c49\u65cf\uff09",
    ),
    ("\u501f\u6b3e\u65e5\u671f", "2025\u5e741\u670810\u65e5"),
    (
        "\u501f\u6b3e\u91d1\u989d",
        "20\u4e07\u5143\uff08\u94f6\u884c\u8f6c\u8d2610\u4e07 + \u652f\u4ed8\u5b9d\u4ee3\u4ed810\u4e07\uff09",
    ),
    (
        "\u539f\u544a\u4e3b\u5f20",
        "\u88ab\u544a\u4ee5\u77ed\u671f\u8d44\u91d1\u5468\u8f6c\u4e3a\u7531\u501f\u6b3e\uff0c\u5e94\u507f\u8fd8\u672c\u91d1+\u5229\u606f",
    ),
    (
        "\u88ab\u544a\u6297\u8fa9",
        "\u6b3e\u9879\u7cfb\u8001\u5e84\u501f\u6b3e\uff0c\u5c0f\u9648\u4ec5\u4e3a\u4ee3\u6536\u4ee3\u4ed8\uff08\u8d26\u6237\u7531\u8001\u5e84\u4f7f\u7528\uff09",
    ),
    (
        "\u6838\u5fc3\u4e89\u8bae",
        "\u501f\u8d37\u5173\u7cfb\u4e3b\u4f53\uff1a\u5c0f\u9648 vs \u8001\u5e84\uff1b\u662f\u5426\u5b58\u5728\u9762\u5bf9\u9762\u501f\u6b3e\u5408\u610f",
    ),
]
table = doc.add_table(rows=len(summary_rows), cols=2)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (k, v) in enumerate(summary_rows):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v
_set_table_font(table)

# ---------------------------------------------------------------------------
# 争点列表 (Heading 1 + Table)
# ---------------------------------------------------------------------------
# Extract issues from result rounds (first round plaintiff output has issue_ids)
# Better: parse from the actual data
all_outputs = []
for r in result.get("rounds", []):
    for o in r.get("outputs", []):
        all_outputs.append(o)

# Collect unique issue_ids from all outputs
issue_ids_seen = []
issue_ids_set = set()
for o in all_outputs:
    for iid in o.get("issue_ids", []):
        if iid not in issue_ids_set:
            issue_ids_seen.append(iid)
            issue_ids_set.add(iid)

# Get conflict issue_ids
conflict_issues = {c["issue_id"] for c in result.get("evidence_conflicts", [])}

n_issues = len(issue_ids_seen)
n_burdens = 8  # from the run log

doc.add_heading(
    "\u4e89\u70b9\u5217\u8868\uff08{}\u4e2a\u4e89\u70b9\uff09".format(n_issues), level=1
)

table2 = doc.add_table(rows=1, cols=4)
table2.style = "Light Grid Accent 1"
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table2.rows[0].cells
hdr[0].text = "#"
hdr[1].text = "\u4e89\u70b9"
hdr[2].text = "\u7c7b\u578b"
hdr[3].text = "\u72b6\u6001"

for idx, iid in enumerate(issue_ids_seen, 1):
    row = table2.add_row().cells
    num = iid.split("-")[-1] if "-" in iid else str(idx)
    row[0].text = num
    desc, itype = _issue_info.get(iid, (iid, ""))
    row[1].text = desc
    row[2].text = _ISSUE_TYPE_ZH.get(itype, itype)
    status = "\u26a0 \u51b2\u7a81" if iid in conflict_issues else "\u2714"
    row[3].text = status
_set_table_font(table2)

# ---------------------------------------------------------------------------
# 三轮对抗记录
# ---------------------------------------------------------------------------
doc.add_heading("\u4e09\u8f6e\u5bf9\u6297\u8bb0\u5f55", level=1)

round_labels = {
    "claim": "\u9996\u8f6e\u4e3b\u5f20",
    "evidence": "\u8bc1\u636e\u6574\u7406",
    "rebuttal": "\u9488\u5bf9\u6027\u53cd\u9a73",
}

for r in result.get("rounds", []):
    rn = r["round_number"]
    phase = r["phase"]
    label = round_labels.get(phase, phase)
    doc.add_heading("第{}轮：{}".format(rn, label), level=2)

    for o in r["outputs"]:
        role = o["agent_role_code"]
        title_text = o.get("title", "")
        body = o.get("body", "")
        ev_cited = ", ".join(o.get("evidence_citations", []))

        # Role label mapping
        if "plaintiff" in role:
            role_label = "[\u539f\u544a\u4ee3\u7406]"
        elif "defendant" in role:
            role_label = "[\u88ab\u544a\u4ee3\u7406]"
        else:
            role_label = "[\u8bc1\u636e\u7ba1\u7406]"

        # Agent title
        p = doc.add_paragraph()
        _add_run(
            p,
            "{} {}".format(role_label, title_text),
            bold=True,
            size=SZ_AGENT_TITLE,
            color=_agent_color(role),
        )

        # Body
        _add_styled_para(
            doc, body[:2000] + ("..." if len(body) > 2000 else ""), size=SZ_BODY, color=CLR_BODY
        )

        # Evidence cited
        _add_styled_para(
            doc, "\u5f15\u7528\u8bc1\u636e: " + ev_cited, size=SZ_EVIDENCE, color=CLR_GRAY
        )

        # Risk flags (including own_weaknesses)
        risk_flags = o.get("risk_flags", [])
        if risk_flags:
            _add_styled_para(
                doc, "\u98ce\u9669\u6807\u8bb0:", bold=True, size=SZ_RISK, color=CLR_ORANGE
            )
            for rf in risk_flags:
                desc = rf.get("description", "")
                flag_id = rf.get("flag_id", "")
                if flag_id.startswith("own-weakness-"):
                    prefix = "\u3010\u672c\u65b9\u8584\u5f31\u70b9\u3011"
                else:
                    prefix = ""
                _add_bullet(doc, prefix + desc, size=SZ_RISK, color=CLR_GRAY)

        doc.add_paragraph()

# ---------------------------------------------------------------------------
# 证据冲突分析
# ---------------------------------------------------------------------------
conflicts = result.get("evidence_conflicts", [])
if conflicts:
    doc.add_heading(
        "\u8bc1\u636e\u51b2\u7a81\u5206\u6790\uff08{}\u6761\uff09".format(len(conflicts)), level=1
    )
    for c in conflicts:
        p = doc.add_paragraph()
        _add_run(p, "[{}] ".format(c["issue_id"]), bold=True, size=SZ_SECTION_HDR, color=CLR_RED)
        doc.add_paragraph()
        _add_styled_para(doc, c["conflict_description"], size=SZ_NORMAL, color=CLR_BODY)

# ---------------------------------------------------------------------------
# LLM综合分析
# ---------------------------------------------------------------------------
summary = result.get("summary")
if summary:
    doc.add_heading("AI综合分析", level=1)
    doc.add_heading("\u6574\u4f53\u6001\u52bf\u8bc4\u4f30", level=2)

    overall = summary.get("overall_assessment", "")
    # Parse structured overall assessment if it contains key sections
    if "plaintiff" in overall.lower() or "\u539f\u544a" in overall:
        # Try to parse structured format
        _add_styled_para(doc, overall, size=SZ_NORMAL, color=CLR_BODY)
    else:
        _add_styled_para(doc, overall, size=SZ_NORMAL, color=CLR_BODY)

    # Plaintiff strongest
    p_args = summary.get("plaintiff_strongest_arguments", [])
    if p_args:
        _add_styled_para(
            doc,
            "\u539f\u544a\u6700\u5f3a\u8bba\u70b9",
            bold=True,
            size=SZ_SECTION_HDR,
            color=CLR_BLUE,
        )
        for a in p_args:
            p = doc.add_paragraph()
            _add_run(p, "[{}] ".format(a["issue_id"]), bold=True, size=SZ_NORMAL, color=CLR_BLUE)
            _add_run(p, a.get("position", ""), size=SZ_NORMAL, color=CLR_BODY)
            reasoning = a.get("reasoning", "")
            if reasoning:
                _add_styled_para(doc, "\u25b6 " + reasoning, size=SZ_NORMAL, color=CLR_GRAY)

    # Defendant strongest
    d_args = summary.get("defendant_strongest_defenses", [])
    if d_args:
        _add_styled_para(
            doc,
            "\u88ab\u544a\u6700\u5f3a\u6297\u8fa9",
            bold=True,
            size=SZ_SECTION_HDR,
            color=CLR_RED,
        )
        for d in d_args:
            p = doc.add_paragraph()
            _add_run(p, "[{}] ".format(d["issue_id"]), bold=True, size=SZ_NORMAL, color=CLR_RED)
            _add_run(p, d.get("position", ""), size=SZ_NORMAL, color=CLR_BODY)
            reasoning = d.get("reasoning", "")
            if reasoning:
                _add_styled_para(doc, "\u25b6 " + reasoning, size=SZ_NORMAL, color=CLR_GRAY)

    # Unresolved issues
    unresolved = summary.get("unresolved_issues", [])
    if unresolved:
        _add_styled_para(
            doc,
            "\u5173\u952e\u672a\u89e3\u51b3\u4e89\u70b9",
            bold=True,
            size=SZ_SECTION_HDR,
            color=CLR_BLUE,
        )
        for u in unresolved:
            iid = u.get("issue_id", "")
            title = u.get("issue_title", "")
            why = u.get("why_unresolved", "")
            _add_bullet(
                doc, "{} {}\uff1a{}".format(iid, title, why), size=SZ_NORMAL, color=CLR_BODY
            )

# ---------------------------------------------------------------------------
# 证据缺失报告
# ---------------------------------------------------------------------------
missing = result.get("missing_evidence_report", [])
if missing:
    doc.add_heading("\u8bc1\u636e\u7f3a\u5931\u62a5\u544a", level=1)
    for m in missing:
        p = doc.add_paragraph()
        party = m.get("missing_for_party_id", "")
        _add_run(
            p,
            "[{}] {}: {}".format(m["issue_id"], party, m["description"]),
            bold=True,
            size=SZ_BODY,
            color=CLR_ORANGE,
        )

# ---------------------------------------------------------------------------
# NEW v3 sections
# ---------------------------------------------------------------------------

# ── 争点影响排序 ──
doc.add_heading("争点影响排序（第三版新增）", level=1)
top5 = exec_summary.get("top5_decisive_issues", [])
_add_styled_para(doc, "前五大决定性争点:", bold=True, size=SZ_SECTION_HDR, color=CLR_BLUE)
for iid in top5:
    _add_bullet(doc, iid, size=SZ_NORMAL, color=CLR_BODY)

# ── 裁判路径树 ──
paths = decision_tree.get("paths", [])
if not paths:
    doc.add_heading("\u88c1\u5224\u8def\u5f84\u6811", level=1)
    _add_styled_para(
        doc,
        "（本次运行未生成裁判路径，可能需重新运行庭后分析流程）",
        size=SZ_NORMAL,
        color=CLR_GRAY,
    )
if paths:
    doc.add_heading(
        "\u88c1\u5224\u8def\u5f84\u6811\uff08{}\u6761\u8def\u5f84\uff09".format(len(paths)), level=1
    )
    for path in paths:
        pid = path.get("path_id", "")
        _add_styled_para(doc, "\u8def\u5f84 " + pid, bold=True, size=SZ_SECTION_HDR, color=CLR_BLUE)

        fields = [
            ("\u89e6\u53d1\u6761\u4ef6", path.get("trigger_condition", "")),
            ("\u89e6\u53d1\u4e89\u70b9", ", ".join(path.get("trigger_issue_ids", []))),
            ("\u5173\u952e\u8bc1\u636e", ", ".join(path.get("key_evidence_ids", []))),
            ("\u53ef\u80fd\u7ed3\u679c", path.get("possible_outcome", "")),
        ]
        ci = path.get("confidence_interval")
        if ci:
            lo = ci.get("lower", 0)
            hi = ci.get("upper", 0)
            fields.append(("\u7f6e\u4fe1\u533a\u95f4", "{:.0%} ~ {:.0%}".format(lo, hi)))
        notes = path.get("path_notes", "")
        if notes:
            fields.append(("\u5907\u6ce8", notes))

        for label, val in fields:
            p = doc.add_paragraph()
            _add_run(p, label + "\uff1a", bold=True, size=SZ_RISK, color=CLR_GRAY)
            _add_run(p, val, size=SZ_RISK, color=CLR_BODY)

        doc.add_paragraph()

    blocking = decision_tree.get("blocking_conditions", [])
    if blocking:
        _add_styled_para(
            doc, "\u963b\u65ad\u6761\u4ef6", bold=True, size=SZ_SECTION_HDR, color=CLR_RED
        )
        for bc in blocking:
            _add_bullet(
                doc,
                "{}: {}".format(bc["condition_id"], bc["description"]),
                size=SZ_NORMAL,
                color=CLR_BODY,
            )

# ── 对方最优攻击链 ──
attacks = attack_chain.get("top_attacks", [])
if not attacks:
    doc.add_heading("\u5bf9\u65b9\u6700\u4f18\u653b\u51fb\u94fe", level=1)
    _add_styled_para(
        doc, "（本次运行未生成攻击链，可能需重新运行庭后分析流程）", size=SZ_NORMAL, color=CLR_GRAY
    )
if attacks:
    doc.add_heading("\u5bf9\u65b9\u6700\u4f18\u653b\u51fb\u94fe", level=1)

    order = attack_chain.get("recommended_order", [])
    p = doc.add_paragraph()
    _add_run(p, "\u653b\u51fb\u65b9\uff1a", bold=True, size=SZ_BODY, color=CLR_RED)
    raw_party = attack_chain.get("owner_party_id", "")
    _add_run(p, _PARTY_ZH.get(raw_party, raw_party), size=SZ_BODY, color=CLR_BODY)
    p2 = doc.add_paragraph()
    _add_run(p2, "\u63a8\u8350\u987a\u5e8f\uff1a", bold=True, size=SZ_BODY, color=CLR_RED)
    _add_run(p2, " \u2192 ".join(order), size=SZ_BODY, color=CLR_BODY)
    doc.add_paragraph()

    for node in attacks:
        nid = node.get("attack_node_id", "")
        _add_styled_para(doc, nid, bold=True, size=SZ_SECTION_HDR, color=CLR_RED)

        attack_fields = [
            ("\u76ee\u6807\u4e89\u70b9", node.get("target_issue_id", "")),
            ("\u653b\u51fb\u8bba\u70b9", node.get("attack_description", "")),
            ("\u6210\u529f\u6761\u4ef6", node.get("success_conditions", "")),
            ("\u652f\u6491\u8bc1\u636e", ", ".join(node.get("supporting_evidence_ids", []))),
            ("\u53cd\u5236\u52a8\u4f5c", node.get("counter_measure", "")),
            ("\u5bf9\u65b9\u8865\u8bc1\u7b56\u7565", node.get("adversary_pivot_strategy", "")),
        ]
        for label, val in attack_fields:
            if not val:
                continue
            p = doc.add_paragraph()
            _add_run(p, label + "：", bold=True, size=SZ_RISK, color=CLR_ORANGE)
            _add_run(p, val, size=SZ_RISK, color=CLR_BODY)
        doc.add_paragraph()

# ── 行动建议 ──
doc.add_heading("\u884c\u52a8\u5efa\u8bae", level=1)

stable_claim = exec_summary.get("current_most_stable_claim", "")
_add_styled_para(doc, "最稳诉请版本：", bold=True, size=SZ_SECTION_HDR, color=CLR_GREEN)
_add_styled_para(
    doc, stable_claim if stable_claim else "（暂无稳定诉请版本）", size=SZ_NORMAL, color=CLR_BODY
)

actions_top3 = exec_summary.get("top3_immediate_actions", [])
if actions_top3 and actions_top3 != "\u672a\u542f\u7528":
    _add_styled_para(doc, "前三项立即行动：", bold=True, size=SZ_SECTION_HDR, color=CLR_ORANGE)
    if isinstance(actions_top3, list):
        for a in actions_top3:
            _add_bullet(doc, a, size=SZ_NORMAL, color=CLR_BODY)
    else:
        _add_styled_para(doc, str(actions_top3), size=SZ_NORMAL, color=CLR_GRAY)

gaps = exec_summary.get("critical_evidence_gaps", [])
if gaps and gaps != "\u672a\u542f\u7528":
    _add_styled_para(
        doc, "\u5173\u952e\u7f3a\u8bc1\uff1a", bold=True, size=SZ_SECTION_HDR, color=CLR_ORANGE
    )
    if isinstance(gaps, list):
        for g in gaps:
            _add_bullet(doc, g, size=SZ_NORMAL, color=CLR_BODY)
    else:
        _add_styled_para(doc, str(gaps), size=SZ_NORMAL, color=CLR_GRAY)

# ── 执行摘要 ──
doc.add_heading("\u6267\u884c\u6458\u8981", level=1)

exec_fields = [
    ("前五大决定性争点", exec_summary.get("top5_decisive_issues", [])),
    ("前三项对方最优攻击", exec_summary.get("top3_adversary_optimal_attacks", [])),
]
for label, val in exec_fields:
    p = doc.add_paragraph()
    _add_run(p, label + "\uff1a", bold=True, size=SZ_SECTION_HDR, color=CLR_BLUE)
    if isinstance(val, list):
        for item in val:
            _add_bullet(doc, item, size=SZ_NORMAL, color=CLR_BODY)
    else:
        _add_styled_para(doc, str(val), size=SZ_NORMAL, color=CLR_BODY)

# Amount check
check = amount_report.get("consistency_check_result", {})
verdict_block = check.get("verdict_block_active", False)
p = doc.add_paragraph()
_add_run(
    p,
    "\u91d1\u989d\u4e00\u81f4\u6027\u6821\u9a8c\uff1a",
    bold=True,
    size=SZ_SECTION_HDR,
    color=CLR_BLUE,
)
_add_run(
    p,
    "阻断裁判={}，未解决冲突={}条".format(
        "是" if verdict_block else "否", len(check.get("unresolved_conflicts", []))
    ),
    size=SZ_NORMAL,
    color=CLR_GREEN if not verdict_block else CLR_RED,
)

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
dest = Path(r"C:\Users\david\Desktop") / "v3_\u5bf9\u6297\u5206\u6790\u62a5\u544a.docx"
doc.save(str(dest))
print("Saved:", dest)
