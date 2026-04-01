#!/usr/bin/env python3
"""
关键词类案检索脚本 — 纯本地，无 LLM 调用。
Keyword-based similar case search — local only, no LLM calls.

用法 / Usage:
    python scripts/run_similar_case_search.py [output_dir]

若不指定 output_dir，自动选择 outputs/ 下最新的目录。
If output_dir is omitted, the most recent directory under outputs/ is used.
"""
from __future__ import annotations

import json
import os
import sys
from pathlib import Path

# ---------------------------------------------------------------------------
# 常量 / Constants
# ---------------------------------------------------------------------------

REPO_ROOT = Path(__file__).parent.parent
DATA_INDEX = REPO_ROOT / "data" / "court_cases_index.json"

# 本案关键词 — 王某诉陈某/庄某民间借贷纠纷
# Keywords for Wang v Chen/Zhuang private-lending dispute
CASE_CAUSE_OF_ACTION = "民间借贷纠纷"

SEARCH_TERMS = [
    "民间借贷",
    "借款合同",
    "借贷合意",
    "借款人主体",
    "代收代付",
    "资金流向",
    "出借人",
    "共同借款人",
    "债务加入",
    "表见代理",
    "账户实际使用人",
]

DISPUTE_FOCUSES = [
    "借款主体认定",
    "借贷关系成立",
    "资金转账性质",
    "代收代付抗辩",
    "借贷合意认定",
]


# ---------------------------------------------------------------------------
# 搜索逻辑 / Search logic
# ---------------------------------------------------------------------------

def load_index(path: Path) -> list[dict]:
    raw = path.read_bytes().decode("utf-8")
    return json.loads(raw)


def score_case(entry: dict, search_terms: list[str], cause_of_action: str) -> float:
    """Score a case entry by keyword overlap — same logic as LocalCaseSearcher."""
    score = 0.0
    entry_coa = entry.get("cause_of_action", "")
    entry_keywords = entry.get("keywords", [])
    entry_summary = entry.get("summary", "")

    # Cause-of-action match
    if entry_coa == cause_of_action:
        score += 3.0
    elif cause_of_action in entry_coa or entry_coa in cause_of_action:
        score += 1.0

    # Keyword / summary matches
    for term in search_terms:
        t = term.lower()
        if any(t in kw.lower() or kw.lower() in t for kw in entry_keywords):
            score += 2.0
        elif t in entry_summary.lower():
            score += 1.0

    return score


def search_similar_cases(
    index: list[dict],
    search_terms: list[str],
    cause_of_action: str,
    max_results: int = 10,
) -> list[tuple[float, dict]]:
    scored = []
    all_terms = search_terms
    for entry in index:
        s = score_case(entry, all_terms, cause_of_action)
        if s > 0:
            scored.append((s, entry))
    scored.sort(key=lambda x: x[0], reverse=True)
    return scored[:max_results]


# ---------------------------------------------------------------------------
# DOCX 插入 / DOCX insertion
# ---------------------------------------------------------------------------

def find_docx(output_dir: Path) -> Path | None:
    for f in output_dir.iterdir():
        if f.suffix == ".docx":
            return f
    return None


def find_latest_output_dir() -> Path:
    outputs = REPO_ROOT / "outputs"
    dirs = sorted(
        [d for d in outputs.iterdir() if d.is_dir() and not d.name.startswith(".")],
        key=lambda d: d.name,
        reverse=True,
    )
    if not dirs:
        raise FileNotFoundError(f"No output directories found in {outputs}")
    return dirs[0]


def insert_similar_cases_section(docx_path: Path, results: list[tuple[float, dict]]) -> None:
    """Insert a 类案检索参考 section before the last paragraph of the DOCX."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import lxml.etree as etree

    doc = Document(str(docx_path))

    # Remove existing 类案检索参考 section if present (idempotent)
    paras_to_remove = []
    inside_section = False
    for p in doc.paragraphs:
        if "类案检索参考" in p.text:
            inside_section = True
        if inside_section:
            paras_to_remove.append(p)
        # Stop at next Heading 1 that is NOT the section itself
        if inside_section and p.style.name == "Heading 1" and "类案检索参考" not in p.text:
            paras_to_remove.pop()  # keep the next heading
            inside_section = False

    for p in paras_to_remove:
        p._element.getparent().remove(p._element)

    # Find insertion point: before the last non-empty paragraph (usually a footer note)
    # or simply append to the end of the body
    body = doc.element.body

    def add_heading(text: str, level: int = 1) -> None:
        p = doc.add_heading(text, level=level)

    def add_paragraph(text: str, bold: bool = False) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        if bold:
            run.bold = True

    def add_hyperlink_paragraph(label: str, url: str) -> None:
        """Add a paragraph with a clickable hyperlink."""
        p = doc.add_paragraph()
        # Add label text
        p.add_run(label)
        # Add hyperlink
        from docx.opc.part import Part
        r_id = doc.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(color)
        rPr.append(u)
        new_run.append(rPr)
        t = OxmlElement("w:t")
        t.text = url
        new_run.append(t)
        hyperlink.append(new_run)
        p._p.append(hyperlink)

    # Add section heading
    add_heading("类案检索参考", level=1)
    add_paragraph(
        f"以下类案通过关键词匹配从人民法院案例库（共 5216 条）中检索得出，"
        f"共返回 {len(results)} 条相关案例，按匹配得分降序排列。",
    )

    for rank, (score, entry) in enumerate(results, start=1):
        case_number = entry.get("case_number", "")
        court = entry.get("court", "")
        coa = entry.get("cause_of_action", "")
        keywords = "、".join(entry.get("keywords", []))
        summary = entry.get("summary", "")
        url = entry.get("url", "")

        # Case heading
        add_paragraph(f"{rank}. {case_number}", bold=True)
        add_paragraph(f"法院：{court}")
        add_paragraph(f"案由：{coa}　关键词：{keywords}")
        add_paragraph(f"匹配得分：{score:.1f}")
        if summary:
            # Truncate very long summaries
            display_summary = summary[:300] + "…" if len(summary) > 300 else summary
            add_paragraph(f"裁判要旨：{display_summary}")
        if url:
            add_hyperlink_paragraph("案例链接：", url)
        # Blank separator
        doc.add_paragraph()

    doc.save(str(docx_path))
    print(f"[OK] 已将类案检索结果插入 DOCX：{docx_path}")


# ---------------------------------------------------------------------------
# 主流程 / Main
# ---------------------------------------------------------------------------

def main() -> None:
    # Determine output directory
    if len(sys.argv) >= 2:
        output_dir = Path(sys.argv[1])
    else:
        output_dir = find_latest_output_dir()

    print(f"[INFO] 输出目录：{output_dir}")

    # Find DOCX
    docx_path = find_docx(output_dir)
    if docx_path is None:
        print(f"[ERROR] 未找到 DOCX 文件于 {output_dir}", file=sys.stderr)
        sys.exit(1)
    print(f"[INFO] DOCX 文件：{docx_path.name}")

    # Load index
    if not DATA_INDEX.exists():
        print(f"[ERROR] 案例索引不存在：{DATA_INDEX}", file=sys.stderr)
        sys.exit(1)
    print(f"[INFO] 加载案例索引：{DATA_INDEX}")
    index = load_index(DATA_INDEX)
    print(f"[INFO] 共 {len(index)} 条案例")

    # Search
    all_terms = SEARCH_TERMS + DISPUTE_FOCUSES
    results = search_similar_cases(index, all_terms, CASE_CAUSE_OF_ACTION, max_results=8)
    print(f"[INFO] 找到 {len(results)} 条匹配案例：")
    for score, entry in results:
        print(f"  {score:.1f}  {entry.get('case_number', '')}  {entry.get('cause_of_action', '')}")

    # Save to JSON alongside DOCX
    json_out = output_dir / "similar_cases_keyword.json"
    with open(json_out, "w", encoding="utf-8") as f:
        json.dump(
            [{"score": score, "case": entry} for score, entry in results],
            f,
            ensure_ascii=False,
            indent=2,
        )
    print(f"[INFO] JSON 结果已保存：{json_out}")

    # Insert into DOCX
    insert_similar_cases_section(docx_path, results)


if __name__ == "__main__":
    main()
